from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_五章送审终版.docx")
OUT = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_学术修订版.docx")


def find_paragraph(doc: Document, startswith: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(startswith)


def set_para_text(doc: Document, startswith: str, text: str) -> None:
    find_paragraph(doc, startswith).text = text


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    para = Paragraph(new_p, paragraph._parent)
    para.style = paragraph.style
    para.text = text
    return para


def replace_all_text(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)


def main():
    doc = Document(str(SRC))

    # Abstracts
    set_para_text(
        doc,
        "本文以小红书平台3403位粉丝数10万以上的生活记录类达人为研究对象",
        "本文以小红书平台3403位粉丝数10万以上的生活记录类达人为研究对象，综合运用描述性统计分析、基于规则的文本分析及 K-means 聚类方法，围绕一个核心研究问题展开分析，即在用户画像视角下，高影响力达人呈现出怎样的行为模式，以及这些行为模式如何在内容表达、互动结构与商业化表现上分化。研究发现如下：第一，高影响力生活类达人整体呈现女性主导、腰部达人为主、独立运营居多的画像结构，核心受众主要集中于高线省市的18至34岁年轻女性，且样本整体表现出相对稳定的粉丝质量与互动质量。第二，达人在内容表达上普遍采用“泛生活+垂直兴趣”的复合布局，并通过日常化、可亲近的人设维系与粉丝的关系；评论互动则主要表现为颜值赞美、拟亲密互动、提问求回应与轻度消费转化。第三，经 K-means 聚类后，样本可被描述为均衡主流型、收藏转化与商业合作型以及高评论互动型三类群体，但群体间差异更多体现为方向性的结构偏移，而非完全分离的强分化。"
    )
    set_para_text(
        doc,
        "本研究从达人供给侧视角丰富了现有文献对高影响力创作者群体的实证认识",
        "本研究从达人供给侧视角补充了对高影响力创作者群体的描述性认识，并尝试将用户画像、文本表现与群体分化纳入同一分析框架，为内容创作者定位、品牌达人筛选以及平台创作者分层提供审慎的经验参考。"
    )
    set_para_text(
        doc,
        "The rise of content-driven e-commerce has reshaped both consumer decision-making and brand communication.",
        "Content-driven e-commerce has reshaped consumer decision-making and brand communication on platforms such as Xiaohongshu. Yet existing research still provides limited systematic description of high-impact lifestyle creators from the supply side. This study examines 3,403 Xiaohongshu lifestyle creators with at least 100,000 followers and combines descriptive statistics, rule-based text analysis, and K-means clustering to examine how their behavioral patterns are expressed through content, interaction, and commercialization. The findings are threefold. First, the sample shows a relatively stable profile structure, with a strong concentration of female and mid-tier creators and a core audience of young women in higher-tier cities. Second, creators generally adopt a hybrid content pattern that combines broad lifestyle expression with selected vertical interests, while maintaining closeness through everyday persona construction and comment interaction. Third, the clustering results suggest three creator groups with different emphases on balance, collection-and-conversion orientation, and comment interaction, although the boundaries between groups remain moderate rather than sharply separated. Overall, the study provides a cautious descriptive account of how user profiles, textual expression, and creator differentiation are related on Xiaohongshu."
    )

    # Research feature wording
    set_para_text(doc, "1.4.2 论文创新点", "1.4.2 本文的研究特色")
    set_para_text(
        doc,
        "本文的创新点主要体现在两个方面。",
        "本文的研究特色主要体现在两个方面。其一，围绕一个核心研究问题，将用户画像、文本分析与群体划分纳入同一解释框架，以识别高影响力生活类达人的差异化表现。其二，在结构化画像数据之外，引入较大规模的帖子与评论文本，并将群体划分结果与文本表现联动分析，从而提升对高影响力达人内容供给、互动组织与商业化路径的综合描述能力。"
    )

    # Analysis framework
    frame_heading = find_paragraph(doc, "2.3.3 本文的分析框架")
    insert_paragraph_after(
        frame_heading,
        "本文的分析框架并不将意见领袖理论和社会认同理论作为结论性的贴标签工具，而是将其作为解释性透镜嵌入后续分析。具体而言，意见领袖理论主要用于理解高影响力达人在账号体量、内容保存价值、讨论互动倾向和商业合作强度上的分层差异；社会认同理论则主要用于理解评论中的拟亲密互动、颜值赞美、求回应表达及其与粉丝年龄结构、关注焦点之间的对应关系。基于此，第三章分别设置画像变量、聚类变量和文本规则，第四章再依次报告画像结构、群体划分和文本表现，从而形成“用户画像基础—群体分化结果—文本外显表现”的分析路径。"
    )

    # Method chapter honesty
    set_para_text(
        doc,
        "在文本处理环节，本文统一完成清洗、中文分词、停用词过滤、TF-IDF辅助筛词、高频词筛选、共现网络构建与评论互动规则识别。",
        "在文本处理环节，本文统一完成清洗、中文分词、停用词过滤、TF-IDF辅助筛词、高频词筛选、共现网络构建与评论互动规则识别。其中，共现关系以同一帖子中的关键词共现为统计窗口，评论互动类型识别采用基于关键词与语义模式的规则归类方式，以保证结果具有较好的可解释性与可复核性，相关规则见表3-3。TF-IDF辅助筛词和共词分析方法分别参考 Salton and Buckley（1988）与 Callon et al.（1983）的相关研究。需要说明的是，本文并未进一步引入主题模型、深层语义模型或人工标注一致性检验，因此这里的“文本分析”更适合被理解为基于规则的解释性文本归纳，而非高阶语义建模。"
    )
    set_para_text(
        doc,
        "为避免群体差异仅停留于描述性比较，本文进一步设置两类轻量统计检验：一是对群体与评论互动类型构成的列联表执行卡方检验，以判断互动结构差异是否具有统计意义；二是对关键连续变量执行 Kruskal-Wallis 检验，以识别不同群体在粉丝质量、互动质量与商业化指标上的中位数差异。",
        "为避免群体差异仅停留于描述性比较，本文进一步设置两类轻量统计检验：一是对群体与评论互动类型构成的列联表执行卡方检验，以判断互动结构差异是否具有统计意义；二是对关键连续变量执行 Kruskal-Wallis 检验，以识别不同群体在粉丝质量、互动质量与商业化指标上的中位数差异。考虑到文本归类本身采用规则法，这些检验的作用主要是为群体差异提供有限的统计支持，而非建立强因果解释。"
    )

    # Clustering honesty
    set_para_text(
        doc,
        "在用户画像视角下，聚类分析的目的不是单纯把达人分组，而是识别不同受众结构下达人在互动风格和商业化表现上的组合差异。",
        "在用户画像视角下，聚类分析的目的不是单纯把达人分组，而是识别不同受众结构下达人在互动风格和商业化表现上的组合差异。本文之所以仍采用 K-means，而未进一步切换至 DBSCAN 或层次聚类，一方面是因为核心变量数量较少且以连续指标为主，另一方面是考虑到 K-means 在结果解释和群体概括上更便于与画像指标和文本分析对接。需要强调的是，这一选择更多服务于描述性分层，而不意味着 K-means 是当前样本的唯一最优方案。"
    )
    set_para_text(
        doc,
        "结果显示，当K=3时，轮廓系数达到0.354，为当前比较范围内的最高值，同时SSE曲线在K=2至K=3之间出现明显拐点。",
        "结果显示，当K=3时，轮廓系数达到0.354，为当前比较范围内的最高值，同时SSE曲线在K=2至K=3之间出现明显拐点。需要说明的是，按照聚类质量评估的一般经验，0.354属于中等偏弱的聚类结构，并不意味着群体边界已经十分清晰。对于生活记录类达人这类同质性较高、边界相对模糊的样本而言，该结果只能说明样本中存在一定程度的分化趋势，而不宜被夸大为强分离的稳定类型。进一步以不同 random_state 重复聚类后，调整兰德指数（ARI）介于 0.9897 至 1.0000 之间，说明在当前变量和模型设定下，三类划分在初始化变化下总体稳定。综合判断，本文采用K=3的聚类方案对3371位达人进行解释性分层。"
    )
    set_para_text(
        doc,
        "基于K-means聚类结果，本文将3371位高影响力达人划分为三类特征鲜明的群体，相关画像特征概览见表4-1。",
        "基于K-means聚类结果，本文将3371位高影响力达人划分为三类具有解释意义的群体，相关画像特征概览见表4-1。群体0共有2671位达人，占比最高，整体表现较为均衡，可概括为“均衡主流型”；群体1共有617位达人，在藏赞比和商业笔记占比上明显高于整体均值，可概括为“收藏转化与商业合作型”；群体2共有83位达人，评赞比显著高于其余两类，可概括为“高评论互动型”。需要再次强调的是，群体2样本量相对较小，因此后续分析更适合作为观察性补充，而不宜被解释为具有强泛化性的稳定主类型。"
    )
    set_para_text(
        doc,
        "收藏转化与商业合作型在商业笔记总数、收藏导向互动以及报价水平上更为突出，显示出较强的内容保存价值与商业承接能力；高评论互动型则在评论导向互动上更为明显，但其商业笔记总数和报价水平相对较低，说明互动热度与商业合作强度并不完全同步。",
        "收藏转化与商业合作型在商业笔记总数、收藏导向互动以及报价水平上更为突出，显示出较强的内容保存价值与商业承接能力；高评论互动型则在评论导向互动上更为明显，但其商业笔记总数和报价水平相对较低，说明互动热度与商业合作强度并不完全同步。由于该群体样本规模有限，下文仅将其作为补充性观察对象，而不与前两类做等量展开。"
    )

    # Comment analysis cooling
    set_para_text(
        doc,
        "为检验上述差异是否仅为描述性观察，本文进一步对群体与评论互动类型构成的列联表进行卡方检验，并对关键连续变量执行 Kruskal-Wallis 检验，结果见表4-3。",
        "为检验上述差异是否仅为描述性观察，本文进一步对群体与评论互动类型构成的列联表进行卡方检验，并对关键连续变量执行 Kruskal-Wallis 检验，结果见表4-3。卡方检验显示，不同群体在评论互动类型分布上存在统计显著差异（χ²=370.94, p<0.001），但 Cramér's V=0.039，说明效应量极弱，群体变量对评论类型分布的实际区分度很小。因此，这里的差异更适合被理解为大样本条件下可识别的方向性偏移，而不应被解读为高度分化的评论结构。"
    )

    # Theory / contribution / practice / limitations
    set_para_text(
        doc,
        "从理论层面看，本文围绕“高影响力达人呈现出怎样的行为模式及其分化”这一核心问题，补充了对意见领袖理论和社会认同理论的描述性理解。",
        "从理论层面看，本文并未试图“验证”意见领袖理论和社会认同理论，而是将二者作为解释当前发现的分析透镜。意见领袖理论帮助我们理解，高影响力达人并非只依赖单一流量优势，而可能通过持续内容供给、较强的保存价值或高参与评论等不同路径形成影响力；社会认同理论则有助于解释评论区中拟亲密互动、颜值赞美与求回应表达为何会与特定受众结构共同出现。需要强调的是，评论互动差异虽然达到统计显著，但效应量总体较弱，因此理论讨论更适合停留在“补充”和“细化”的层面，而非得出强因果或强验证结论。"
    )
    set_para_text(
        doc,
        "从实践层面看，达人自身不宜只围绕粉丝规模理解影响力，而应根据自身更接近的行为模式路径优化策略：均衡主流型更应保持稳定的泛生活内容供给与受众黏性，收藏转化与商业合作型可进一步强化内容的保存价值与商业承接能力，高评论互动型则需在维持互动热度的同时审慎推进商业合作。",
        "从实践层面看，达人自身不宜只围绕粉丝规模理解影响力，而应根据自身更接近的模式路径优化策略：均衡主流型更应保持稳定的泛生活内容供给与受众黏性，收藏转化与商业合作型可进一步强化内容的保存价值与商业承接能力。对于高评论互动型，由于样本规模较小且商业化表现并不稳定，本文更倾向于将其视为补充性观察，不建议据此直接提炼强操作性策略。"
    )
    set_para_text(
        doc,
        "尽管本文在数据规模与分析维度上具有一定系统性，但仍存在若干局限。",
        "尽管本文在数据规模与分析维度上具有一定系统性，但仍存在若干局限。首先，本文所讨论的“行为模式”并非严格意义上的时间序列行为，而是达人在内容表达、互动方式与商业化表现上的稳定组合，因此结论更接近结构性特征描述，而非对动态行为过程的还原。"
    )
    set_para_text(
        doc,
        "再次，结构化画像数据主要来自第三方平台整理结果，仍可能存在覆盖范围与口径选择上的平台偏差；评论与帖子文本虽然规模较大，但采集主要集中于论文研究阶段的单一时间窗口，仍受到抓取时点和可见范围的限制。",
        "再次，结构化画像数据主要来自第三方平台整理结果，仍可能存在覆盖范围与口径选择上的平台偏差；评论与帖子文本虽然规模较大，但采集主要集中于论文研究阶段的单一时间窗口，仍受到抓取时点和可见范围的限制。最后，本文主要采用描述性统计、基于规则的文本分析与聚类分析方法，并补充了轻量统计检验，但尚未进一步引入回归模型、人工标注一致性验证或更复杂的主题模型，因此当前结论更适合被理解为对高影响力达人群体特征及其差异的描述性证据。"
    )

    # Figure numbering and references
    replacements = {
        "图5 图文与视频内容形式效能对比": "图4-1 图文与视频内容形式效能对比",
        "图6 肘部法则与轮廓系数K值选取": "图4-2 肘部法则与轮廓系数K值选取",
        "图1呈现了达人标签的共现网络结构。": "图4-3呈现了达人标签的共现网络结构。",
        "图2展示了3403位达人的粉丝Top1关注焦点分布。": "图4-4展示了3403位达人的粉丝Top1关注焦点分布。",
        "结合图1和图2可以进一步发现": "结合图4-3和图4-4可以进一步发现",
        "图1 达人标签共现词网络": "图4-3 达人标签共现词网络",
        "图2 高影响力达人粉丝关注焦点分布": "图4-4 高影响力达人粉丝关注焦点分布",
        "图3为3403位达人个人简介的高频词词云。": "图4-5为3403位达人个人简介的高频词词云。",
        "图3 高影响力达人简介高频词": "图4-5 高影响力达人简介高频词",
        "图4中的时段分布": "图4-6中的时段分布",
        "图4 高影响力达人粉丝高峰活跃时段与星期分布": "图4-6 高影响力达人粉丝高峰活跃时段与星期分布",
        "图6给出了不同 K 值下 SSE 与轮廓系数的变化情况。": "图4-2给出了不同 K 值下 SSE 与轮廓系数的变化情况。"
    }
    for old, new in replacements.items():
        replace_all_text(doc, old, new)

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
