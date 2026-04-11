from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from importlib.util import module_from_spec, spec_from_file_location
from scipy.stats import chi2_contingency, kruskal
from sklearn.cluster import KMeans
from sklearn.metrics import adjusted_rand_score
from sklearn.preprocessing import StandardScaler


BASE_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_单问题重构版.docx")
OUT_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_五章统计增强版.docx")
BASE_DIR = Path("/Users/dxk/xiaohongshu")
TABLE_OUT = BASE_DIR / "reports" / "chapter4_support" / "tables" / "table_4_3_group_stat_tests.csv"


def load_helpers():
    p = BASE_DIR / "analysis" / "11_chapter4_support.py"
    spec = spec_from_file_location("c4_support", p)
    mod = module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            yield ("t", Table(child, doc))


def paragraph_block_index(doc, startswith: str) -> int:
    for idx, (kind, obj) in enumerate(iter_block_items(doc)):
        if kind == "p" and obj.text.strip().startswith(startswith):
            return idx
    raise ValueError(startswith)


def find_paragraph(doc, startswith: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(startswith)


def find_paragraph_with_next(doc, text: str, next_startswith: str) -> Paragraph:
    paras = doc.paragraphs
    for i, p in enumerate(paras[:-1]):
        if p.text.strip() == text:
            for j in range(i + 1, len(paras)):
                nxt = paras[j].text.strip()
                if nxt:
                    if nxt.startswith(next_startswith):
                        return p
                    break
    raise ValueError(f"{text} -> {next_startswith}")


def set_para_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = text


def remove_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def insert_paragraph_before(paragraph: Paragraph, text: str, style) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.text = text
    return new_para


def insert_paragraph_after(paragraph: Paragraph, text: str, style) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.text = text
    return new_para


def move_block_range_before(doc, start_text: str, end_text: str, before_text: str) -> None:
    body = doc.element.body
    blocks = list(body.iterchildren())
    start = paragraph_block_index(doc, start_text)
    end = paragraph_block_index(doc, end_text)
    before = paragraph_block_index(doc, before_text)
    segment = blocks[start:end]
    for el in segment:
        body.remove(el)
    before = paragraph_block_index(doc, before_text)
    for el in reversed(segment):
        body.insert(before, el)


def move_caption_and_table_before(doc, caption_text: str, before_text: str) -> None:
    body = doc.element.body
    blocks = list(body.iterchildren())
    cap_idx = paragraph_block_index(doc, caption_text)
    if cap_idx + 1 >= len(blocks):
        raise ValueError(f"Missing table after caption {caption_text}")
    segment = blocks[cap_idx:cap_idx + 2]
    for el in segment:
        body.remove(el)
    before_idx = paragraph_block_index(doc, before_text)
    for el in reversed(segment):
        body.insert(before_idx, el)


def remove_paragraphs_starting_with(doc, prefixes: list[str]) -> None:
    for prefix in prefixes:
        for p in list(doc.paragraphs):
            if p.text.strip().startswith(prefix):
                remove_paragraph(p)
                break


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_table_three_line(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ["left", "right", "insideH", "insideV"]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")
    for edge, size in [("top", "12"), ("bottom", "12")]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"})


def insert_table_after(paragraph: Paragraph, rows: list[list[str]]):
    doc = paragraph._parent
    table = doc.add_table(rows=1, cols=len(rows[0]), width=Inches(6.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[j], text, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, text in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], text, bold=False, align=align)
    set_table_three_line(table)
    paragraph._p.addnext(table._tbl)
    return table


def compute_stats():
    mod = load_helpers()
    clusters = pd.read_csv(BASE_DIR / "data" / "daren_clusters_k3.csv")
    clusters["creator_id"] = clusters["达人官方地址"].astype(str).str.extract(r"profile/([a-f0-9]{24})")
    comments = pd.read_csv(BASE_DIR / "data" / "dom_crawl" / "comments.csv")
    comments["creator_id"] = comments["creator_id"].astype(str)
    comments = comments.merge(clusters[["creator_id", "群体标签_k3"]], on="creator_id", how="inner")
    comments["content"] = comments["content"].fillna("").map(mod.clean_text)
    comments = comments[comments["content"].str.len() >= 2].copy()
    comments["interaction_label"] = comments["content"].map(mod.label_comment)
    ctab = pd.crosstab(comments["群体标签_k3"], comments["interaction_label"])
    chi2, p_value, dof, _ = chi2_contingency(ctab)
    n = ctab.values.sum()
    cramers_v = (chi2 / (n * (min(ctab.shape[0] - 1, ctab.shape[1] - 1)))) ** 0.5

    features = ["粉丝数", "Top1年龄段占比", "藏赞比", "评赞比", "商业笔记占比"]
    X = clusters[features].apply(pd.to_numeric, errors="coerce")
    Xs = StandardScaler().fit_transform(X)
    base_labels = KMeans(n_clusters=3, init="k-means++", n_init=10, random_state=42).fit_predict(Xs)
    ari_values = []
    for seed in [0, 1, 7, 21, 42, 99, 2026]:
        labels = KMeans(n_clusters=3, init="k-means++", n_init=10, random_state=seed).fit_predict(Xs)
        ari_values.append(adjusted_rand_score(base_labels, labels))
    ari_min = min(ari_values)
    ari_max = max(ari_values)

    kruskal_vars = [
        "活跃粉丝占比",
        "水粉占比",
        "近60天平均评论",
        "商业笔记总数",
        "图文笔记报价",
        "视频笔记报价",
    ]
    rows = []
    name_map = {
        "活跃粉丝占比": "活跃粉丝占比",
        "水粉占比": "水粉占比",
        "近60天平均评论": "近60天平均评论",
        "商业笔记总数": "商业笔记总数",
        "图文笔记报价": "图文笔记报价",
        "视频笔记报价": "视频笔记报价",
    }
    for var in kruskal_vars:
        groups = []
        medians = []
        for g in sorted(clusters["群体标签_k3"].dropna().unique()):
            s = pd.to_numeric(clusters.loc[clusters["群体标签_k3"] == g, var], errors="coerce").dropna()
            groups.append(s)
            medians.append(s.median())
        stat, p = kruskal(*groups)
        rows.append({
            "检验对象": name_map[var],
            "检验方法": "Kruskal-Wallis",
            "统计量": f"H={stat:.2f}",
            "p值": "<0.001" if p < 0.001 else f"{p:.3f}",
            "结果解释": "差异显著" if p < 0.05 else "差异不显著",
            "群体0中位数": medians[0],
            "群体1中位数": medians[1],
            "群体2中位数": medians[2],
        })

    rows.insert(0, {
        "检验对象": "评论互动类型分布",
        "检验方法": "卡方检验",
        "统计量": f"χ²={chi2:.2f}; V={cramers_v:.3f}",
        "p值": "<0.001" if p_value < 0.001 else f"{p_value:.3f}",
        "结果解释": "差异显著，但效应量较弱",
        "群体0中位数": None,
        "群体1中位数": None,
        "群体2中位数": None,
    })
    out_df = pd.DataFrame(rows)
    TABLE_OUT.parent.mkdir(parents=True, exist_ok=True)
    out_df.to_csv(TABLE_OUT, index=False, encoding="utf-8-sig")

    return {
        "chi2": chi2,
        "p_value": p_value,
        "dof": dof,
        "cramers_v": cramers_v,
        "ari_min": ari_min,
        "ari_max": ari_max,
        "stats_df": out_df,
    }


def main():
    stats = compute_stats()
    doc = Document(str(BASE_DOC))

    chapter_style = find_paragraph(doc, "第三章 研究设计与实证分析").style
    section_style = find_paragraph(doc, "3.1 研究设计与数据来源").style
    body_style = find_paragraph(doc, "本文以小红书平台中的高影响力生活记录类达人为研究对象").style
    caption_style = find_paragraph(doc, "表3-1 数据来源与样本规模").style

    # Chapter restructuring.
    set_para_text(find_paragraph(doc, "第三章 研究设计与实证分析"), "第三章 研究设计")
    para_var = find_paragraph(doc, "在指标体系上，本文从基础身份")
    insert_paragraph_before(para_var, "3.2 指标体系与变量设定", section_style)
    para_method = find_paragraph(doc, "在文本处理环节，本文统一完成清洗")
    insert_paragraph_before(para_method, "3.3 文本挖掘、聚类与统计检验设计", section_style)

    set_para_text(find_paragraph(doc, "3.1 研究设计与数据来源"), "3.1 研究对象与数据来源")
    set_para_text(find_paragraph(doc, "本文以小红书平台中的高影响力生活记录类达人为研究对象"), "本文以小红书平台中的高影响力生活记录类达人为研究对象，并以粉丝数达到10万作为进入样本池的操作性门槛。该标准主要用于界定研究范围，而非将高影响力简单等同于单一流量指标；后续分析仍结合互动表现、商业笔记占比与文本特征综合考察达人差异。结构化画像样本来自灰豚平台在论文研究阶段导出的达人总表，共3403位达人；经缺失值处理与异常值筛查后，最终用于K-means聚类的有效样本为3371位。")
    set_para_text(find_paragraph(doc, "本章实际使用的达人画像样本为3403位达人"), "本文实际使用的达人画像样本为3403位达人、聚类样本为3371位达人、文本样本为9731条帖子与123357条评论，对齐后的群体比较样本为9711条帖子与123176条评论，具体见表3-1。")
    set_para_text(find_paragraph(doc, "在数据对齐环节，本文将结构化样本与DOM抓取的帖子、评论数据按照creator_id进行映射"), "在数据对齐环节，本文将结构化样本与DOM抓取的帖子、评论数据按照 creator_id 进行映射，并在必要时进行群体标签回填，从而得到可用于群体比较的文本样本。")

    set_para_text(find_paragraph(doc, "在指标体系上，本文从基础身份"), "在指标体系上，本文从基础身份、互动表现、商业化水平和粉丝画像四个维度构建结构化变量，并将帖子标题、正文、简介与评论文本作为文本变量的主要来源。由于本文关注的是用户画像约束下的行为模式差异，因此用户画像与互动结构被置于解释框架的核心位置。")
    set_para_text(find_paragraph(doc, "本文所称“行为模式”，是指达人在特定受众结构约束下所呈现出的内容表达"), "本文所称“行为模式”，是指达人在特定受众结构约束下所呈现出的内容表达、互动方式与商业化表现的稳定组合。相较于泛泛讨论单一行为动作，本文更关注达人如何在相对稳定的受众基础上形成内容供给风格、评论互动路径与商业合作方式。")
    set_para_text(find_paragraph(doc, "在聚类变量选择上，本文优先保留能够同时反映账号体量"), "在聚类变量选择上，本文优先保留能够同时反映账号体量、受众年龄结构、内容保存价值、讨论互动倾向与商业合作强度的指标，即粉丝数、Top1年龄段占比、藏赞比、评赞比和商业笔记占比。性别比例、地域分布、活跃粉丝占比、水粉占比以及报价、CPE、CPM等变量虽然具有解释价值，但更适合作为群体结果比较变量，而不直接作为决定分组边界的入模变量，相关定义见表3-2。")

    set_para_text(find_paragraph(doc, "在文本处理环节，本文统一完成清洗"), "在文本处理环节，本文统一完成清洗、中文分词、停用词过滤、TF-IDF辅助筛词、高频词筛选、共现网络构建与评论互动规则识别。其中，共现关系以同一帖子中的关键词共现为统计窗口，评论互动类型识别采用规则归类方式，以保证结果具有较好的可解释性与可复核性，相关规则见表3-3。TF-IDF辅助筛词和共词分析方法分别参考 Salton and Buckley（1988）与 Callon et al.（1983）的相关研究。")
    set_para_text(find_paragraph(doc, "在聚类实现上，本文使用 sklearn 的 KMeans 模型"), "在聚类实现上，本文使用 sklearn 的 KMeans 模型，初始化方式为 k-means++，设置 n_init=10、random_state=42，并参考 MacQueen（1967）以及 Arthur and Vassilvitskii（2007）关于 K-means 与初始化改进的讨论。剔除异常值后，再结合肘部法则和轮廓系数确定最终聚类数。")
    para_stats_design = insert_paragraph_after(find_paragraph(doc, "在聚类实现上，本文使用 sklearn 的 KMeans 模型"), "为避免群体差异仅停留于描述性比较，本文进一步设置两类轻量统计检验：一是对群体与评论互动类型构成的列联表执行卡方检验，以判断互动结构差异是否具有统计意义；二是对关键连续变量执行 Kruskal-Wallis 检验，以识别不同群体在粉丝质量、互动质量与商业化指标上的中位数差异。", body_style)
    set_para_text(find_paragraph(doc, "本章首先展示高影响力达人的整体画像结构"), "基于上述研究设计，第四章将依次呈现高影响力达人的画像结构、群体划分结果以及文本表现差异，并在必要处补充轻量统计检验，以增强群体比较的说服力。")

    # Rename design tables before moving.
    set_para_text(find_paragraph(doc, "表3-3 聚类变量定义与入模理由"), "表3-2 聚类变量定义与入模理由")
    set_para_text(find_paragraph(doc, "表3-2 文本预处理与评论互动识别规则"), "表3-3 文本预处理与评论互动识别规则")

    # Insert chapter 4 heading before old 3.2 results heading.
    para_old_32 = find_paragraph(doc, "3.2 高影响力达人画像结构分析")
    chap4 = insert_paragraph_before(para_old_32, "第四章 实证结果分析", chapter_style)
    insert_paragraph_after(chap4, "本章在第三章研究设计的基础上，依次报告画像结构分析、群体划分结果以及文本表现差异，并将关键群体差异通过轻量统计检验进一步验证。", body_style)

    # Renumber empirical section headings.
    set_para_text(find_paragraph(doc, "3.2 高影响力达人画像结构分析"), "4.1 高影响力达人画像结构分析")
    set_para_text(find_paragraph(doc, "3.3 基于 K-means 算法的达人群体划分与类型特征"), "4.2 高影响力达人群体划分与类型特征")
    set_para_text(find_paragraph(doc, "3.3.1 聚类特征变量的选择与模型设定"), "4.2.1 聚类特征变量的选择与模型设定")
    set_para_text(find_paragraph(doc, "3.3.2 三类高影响力达人群体的画像特征解析"), "4.2.2 三类高影响力达人群体的画像特征解析")
    set_para_text(find_paragraph(doc, "3.4 高影响力达人行为模式的文本表现"), "4.3 高影响力达人行为模式的文本表现")
    set_para_text(find_paragraph(doc, "3.4.1 泛生活与垂直兴趣交织的内容生态网络解析"), "4.3.1 泛生活与垂直兴趣交织的内容生态网络解析")
    set_para_text(find_paragraph(doc, "3.4.2 信任机制下的人设与评论互动文本解码"), "4.3.2 信任机制下的人设与评论互动文本解码")
    set_para_text(find_paragraph(doc, "3.4.3 内容发布时机"), "4.3.3 内容发布时机")
    set_para_text(find_paragraph(doc, "3.4.4 内容形式选择"), "4.3.4 内容形式选择")

    # Renumber empirical tables.
    set_para_text(find_paragraph(doc, "表3-4 三类群体画像特征概览"), "表4-1 三类群体画像特征概览")
    set_para_text(find_paragraph(doc, "表3-5 三类群体粉丝质量、互动质量与商业化特征比较"), "表4-2 三类群体粉丝质量、互动质量与商业化特征比较")
    for p in doc.paragraphs:
        if not p.text:
            continue
        p.text = (
            p.text.replace("表3-4", "表4-1")
                  .replace("表3-5", "表4-2")
                  .replace("表3-3", "表3-2")
                  .replace("表3-2", "表3-3_TMP")
        )
    for p in doc.paragraphs:
        if "表3-3_TMP" in p.text:
            p.text = p.text.replace("表3-3_TMP", "表3-3")

    # Repair specific references after chained replacement.
    set_para_text(find_paragraph(doc, "在聚类变量选择上，本文优先保留能够同时反映账号体量"), "在聚类变量选择上，本文优先保留能够同时反映账号体量、受众年龄结构、内容保存价值、讨论互动倾向与商业合作强度的指标，即粉丝数、Top1年龄段占比、藏赞比、评赞比和商业笔记占比。性别比例、地域分布、活跃粉丝占比、水粉占比以及报价、CPE、CPM等变量虽然具有解释价值，但更适合作为群体结果比较变量，而不直接作为决定分组边界的入模变量，相关定义见表3-2。")
    set_para_text(find_paragraph(doc, "在文本处理环节，本文统一完成清洗"), "在文本处理环节，本文统一完成清洗、中文分词、停用词过滤、TF-IDF辅助筛词、高频词筛选、共现网络构建与评论互动规则识别。其中，共现关系以同一帖子中的关键词共现为统计窗口，评论互动类型识别采用规则归类方式，以保证结果具有较好的可解释性与可复核性，相关规则见表3-3。TF-IDF辅助筛词和共词分析方法分别参考 Salton and Buckley（1988）与 Callon et al.（1983）的相关研究。")

    # Strengthen 4.2 with robustness and caution.
    set_para_text(find_paragraph(doc, "为确定最优聚类数K，本文综合使用肘部法则与轮廓系数"), "为确定最优聚类数K，本文综合使用肘部法则与轮廓系数，对K=2至K=8进行逐一比较。聚类实现采用 sklearn 的 KMeans，初始化方式为 k-means++，并设置 n_init=10、random_state=42，以降低随机初始化带来的扰动。图6给出了不同 K 值下 SSE 与轮廓系数的变化情况。")
    set_para_text(find_paragraph(doc, "结果显示，当K=3时，轮廓系数达到0.354"), f"结果显示，当K=3时，轮廓系数达到0.354，为当前比较范围内的最高值，同时SSE曲线在K=2至K=3之间出现明显拐点。需要说明的是，0.354并不属于边界非常清晰的高分聚类结构，但对于生活记录类达人这类同质性较高、边界相对模糊的样本而言，仍可视为具有一定区分度并具备解释价值。进一步以 random_state=0、1、7、21、99 和 2026 进行重复聚类，与基准方案相比的调整兰德指数（ARI）介于 {stats['ari_min']:.4f} 至 {stats['ari_max']:.4f} 之间，说明三类群体的划分在不同初始化条件下总体稳定。综合判断，本文采用K=3的聚类方案对3371位达人进行群体划分。")
    set_para_text(find_paragraph(doc, "基于K-means聚类结果，本文将3371位高影响力达人划分为三类特征鲜明的群体"), "基于K-means聚类结果，本文将3371位高影响力达人划分为三类特征鲜明的群体，相关画像特征概览见表4-1。群体0共有2671位达人，占比最高，整体表现较为均衡，可概括为“均衡主流型”；群体1共有617位达人，在藏赞比和商业笔记占比上明显高于整体均值，可概括为“收藏转化与商业合作型”；群体2共有83位达人，评赞比显著高于其余两类，可概括为“高评论互动型”。需要说明的是，群体2样本量相对较小，因此后续文本比较更适合作为结构性特征观察，而不宜被解释为具有强泛化性的单独类型。")
    set_para_text(find_paragraph(doc, "三类群体的并存表明，小红书高影响力达人生态并非均质结构"), "三类群体的并存表明，小红书高影响力达人生态并非均质结构，而是存在明显的功能分化。进一步结合粉丝质量、互动质量与商业化指标，可以更清楚地识别三类群体在受众基础和变现路径上的差异，详见表4-2。")

    # Add statistical results into 4.3.2
    p_res = find_paragraph(doc, "结果显示，三类群体的评论都以泛互动类表达为主")
    stat_intro = insert_paragraph_after(p_res, f"为检验上述差异是否仅为描述性观察，本文进一步对群体与评论互动类型构成的列联表进行卡方检验，并对关键连续变量执行 Kruskal-Wallis 检验，结果见表4-3。卡方检验显示，不同群体在评论互动类型分布上存在统计显著差异（χ²={stats['chi2']:.2f}, p<0.001），但 Cramér's V={stats['cramers_v']:.3f}，说明效应量较弱，差异更多体现为大样本下的结构偏移而非截然分化。", body_style)
    caption = insert_paragraph_after(stat_intro, "表4-3 群体差异统计检验结果", caption_style)

    stats_rows = [["检验对象", "检验方法", "统计量", "p值", "结果解释"]]
    for _, row in stats["stats_df"].iterrows():
        stats_rows.append([str(row["检验对象"]), str(row["检验方法"]), str(row["统计量"]), str(row["p值"]), str(row["结果解释"])])
    insert_table_after(caption, stats_rows)
    after_table = insert_paragraph_after(caption, "", body_style)
    # table inserted immediately after caption; place summary paragraph after table by finding caption again later
    # workaround: find caption and then use following table's _tbl to add paragraph after
    cap = find_paragraph(doc, "表4-3 群体差异统计检验结果")
    tbl = None
    next_el = cap._p.getnext()
    if next_el is not None and next_el.tag == qn("w:tbl"):
        new_p = OxmlElement("w:p")
        next_el.addnext(new_p)
        summary_para = Paragraph(new_p, cap._parent)
        summary_para.style = body_style
        summary_para.text = "进一步从连续变量差异看，近60天平均评论、商业笔记总数以及图文/视频报价在三类群体之间均表现出统计显著差异，而活跃粉丝占比和水粉占比差异不显著。这意味着三类群体的主要差异更集中在互动表现和商业化承接上，而非粉丝质量本身存在显著断裂。"
    try:
        remove_paragraph(after_table)
    except Exception:
        pass

    # Chapter 5 restructuring.
    set_para_text(find_paragraph(doc, "第四章 结论与展望"), "第五章 讨论与结论")
    set_para_text(find_paragraph(doc, "4.1 研究结论"), "5.1 主要研究发现")
    set_para_text(find_paragraph(doc, "4.2 研究局限性"), "5.5 研究局限与未来展望")
    set_para_text(find_paragraph(doc, "4.3 理论与实践启示"), "5.2 理论讨论")
    set_para_text(find_paragraph(doc, "从理论层面看，本文围绕“高影响力达人呈现出怎样的行为模式及其分化”这一核心问题"), "从理论层面看，本文围绕“高影响力达人呈现出怎样的行为模式及其分化”这一核心问题，补充了对意见领袖理论和社会认同理论的描述性理解。意见领袖理论通常强调具有较高可见度和说服力的节点在传播链路中的作用，而本文发现高影响力达人并不存在单一路径：均衡主流型更接近依靠持续内容供给维持广泛影响力的典型路径，收藏转化与商业合作型则表现出更强的保存价值与商业承接能力，高评论互动型则更多依赖高参与评论和关系表达来维系影响力。由此可见，高影响力并不只体现为单一流量优势，而是对应不同的受众连接方式与影响力实现机制。社会认同理论则有助于理解评论区中的拟亲密互动、颜值赞美与求回应表达。本文发现，评论互动差异虽然具有统计显著性，但效应量总体较弱，这表明社会认同并非只在单一群体中出现，而是以不同强度分布在多个高影响力达人群体之中。")

    # Insert 5.3 and 5.4 after theory discussion.
    theory_para = find_paragraph(doc, "从理论层面看，本文围绕“高影响力达人呈现出怎样的行为模式及其分化”这一核心问题")
    h53 = insert_paragraph_after(theory_para, "5.3 文献回应与研究贡献", section_style)
    p53 = insert_paragraph_after(h53, "与第二章文献述评相比，本文的推进主要体现在三个方面：其一，在研究对象上，从普通用户、单次营销效果或单条内容分析转向高影响力达人供给侧，补充了高影响力创作者群体本身的系统画像；其二，在分析框架上，将用户画像、群体划分与文本表现纳入同一解释路径，避免将达人视为同质化的营销节点；其三，在证据形式上，通过结构化画像数据与帖子、评论文本的对齐分析，进一步展示了不同类型达人在内容表达、互动结构与商业化承接上的差异，从而为现有文献提供了一种更接近创作者分层现实的描述性证据。", body_style)
    h54 = insert_paragraph_after(p53, "5.4 实践启示", section_style)
    p54 = insert_paragraph_after(h54, "从实践层面看，达人自身不宜只围绕粉丝规模理解影响力，而应根据自身更接近的行为模式路径优化策略：均衡主流型更应保持稳定的泛生活内容供给与受众黏性，收藏转化与商业合作型可进一步强化内容的保存价值与商业承接能力，高评论互动型则需在维持互动热度的同时审慎推进商业合作。对于品牌方而言，达人筛选不宜只看粉丝规模和单一报价，而应结合投放目标进行匹配：若强调广泛触达，可优先考虑均衡主流型；若强调内容保存与消费转化，可优先考虑收藏转化与商业合作型；若强调互动热度与评论参与，则可考虑高评论互动型。对于平台而言，研究结果说明创作者分层不宜仅依据单一流量指标，而应结合互动质量、商业承接能力和受众结构开展更精细的治理与扶持。", body_style)

    # Move limitations to the end.
    move_block_range_before(doc, "5.5 研究局限与未来展望", "5.2 理论讨论", "参考文献")
    set_para_text(find_paragraph(doc, "尽管本文在数据规模与分析维度上具有一定系统性"), "尽管本文在数据规模与分析维度上具有一定系统性，但仍存在若干局限。首先，本文使用的是横截面数据，只能反映达人在特定时间节点的静态状态，无法捕捉其从起步到高影响力阶段的动态演变过程，也无法进一步识别变量变化与后续表现之间的因果关系。其次，研究对象聚焦于小红书单一平台，所得结论是否适用于抖音、快手、微博等其他内容平台，仍需结合不同平台的算法机制、用户结构与内容生态进行进一步验证。")
    set_para_text(find_paragraph(doc, "再次，结构化画像数据主要来自第三方平台整理结果"), "再次，结构化画像数据主要来自第三方平台整理结果，仍可能存在覆盖范围与口径选择上的平台偏差；评论与帖子文本虽然规模较大，但采集主要集中于论文研究阶段的单一时间窗口，仍受到抓取时点和可见范围的限制。最后，本文主要采用描述性统计、文本挖掘与聚类分析方法，并补充了轻量统计检验，但尚未进一步引入回归模型、人工标注一致性验证或更复杂的主题模型，因此当前结论更适合被理解为对高影响力达人群体特征及其差异的描述性证据。未来研究可在更长时间窗口下引入纵向数据，并结合更系统的文本模型与因果识别设计，进一步检验不同类型达人行为模式的稳定性与演化路径。")

    # Structural cleanup for chapter 3, chapter 4 tables, and chapter 5 order.
    remove_paragraphs_starting_with(doc, [
        "在聚类实现上，本文使用 sklearn 的 KMeans 模型，初始化方式为 k-means++",
        "为避免群体差异仅停留于描述性比较，本文进一步设置两类轻量统计检验",
        "基于上述研究设计，第四章将依次呈现高影响力达人的画像结构",
        "本文实际使用的达人画像样本为3403位达人、聚类样本为3371位达人",
        "在数据对齐环节，本文将结构化样本与DOM抓取的帖子、评论数据按照 creator_id",
        "之所以未将性别比例、地域分布、活跃粉丝占比、水粉占比以及报价、CPE、CPM直接纳入聚类",
        "在异常值处理方面，本文重点检查了评赞比和商业笔记占比等波动较大的指标",
        "在完成群体划分的基础上，本文进一步从帖子与评论文本出发",
        "具体处理流程包括文本清洗、中文分词、停用词过滤、TF-IDF辅助筛词、高频词筛选、共现网络构建以及互动类型归纳",
    ])

    p31 = find_paragraph(doc, "本文以小红书平台中的高影响力生活记录类达人为研究对象")
    p31b = insert_paragraph_after(p31, "本文实际使用的达人画像样本为3403位达人、聚类样本为3371位达人、文本样本为9731条帖子与123357条评论，对齐后的群体比较样本为9711条帖子与123176条评论，具体见表3-1。", body_style)
    insert_paragraph_after(p31b, "在数据对齐环节，本文将结构化样本与DOM抓取的帖子、评论数据按照 creator_id 进行映射，并在必要时进行群体标签回填，从而得到可用于群体比较的文本样本。", body_style)
    set_para_text(find_paragraph_with_next(doc, "表3-3 聚类变量定义与入模理由", "为确定最优聚类数K"), "表3-2 聚类变量定义与入模理由")
    set_para_text(find_paragraph_with_next(doc, "表3-3 文本预处理与评论互动识别规则", "4.3.1 泛生活与垂直兴趣交织的内容生态网络解析"), "表3-3 文本预处理与评论互动识别规则")
    move_caption_and_table_before(doc, "表3-1 数据来源与样本规模", "3.2 指标体系与变量设定")
    move_caption_and_table_before(doc, "表3-2 聚类变量定义与入模理由", "3.3 文本挖掘、聚类与统计检验设计")
    move_caption_and_table_before(doc, "表3-3 文本预处理与评论互动识别规则", "第四章 实证结果分析")

    # Make chapter 4 flow cleaner.
    move_block_range_before(doc, "从粉丝质量与互动质量看，高影响力达人样本整体表现出较为稳定的受众基础", "4.3 高影响力达人行为模式的文本表现", "4.2 高影响力达人群体划分与类型特征")
    move_caption_and_table_before(doc, "表4-2 三类群体粉丝质量、互动质量与商业化特征比较", "4.3 高影响力达人行为模式的文本表现")

    # Fix chapter 5 ordering.
    move_block_range_before(doc, "其次，从群体划分看，经肘部法则与轮廓系数验证", "5.5 研究局限与未来展望", "5.2 理论讨论")
    move_block_range_before(doc, "5.4 实践启示", "参考文献", "5.5 研究局限与未来展望")
    move_block_range_before(doc, "尽管本文在数据规模与分析维度上具有一定系统性", "5.2 理论讨论", "参考文献")

    # Save updated statistical CSV and doc.
    doc.save(str(OUT_DOC))


if __name__ == "__main__":
    main()
