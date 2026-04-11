#!/usr/bin/env python3
"""
Enhance thesis with fan-quality, engagement-quality, and commercialization fields.
"""

from __future__ import annotations

import csv
import statistics
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


INPUT_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_送审收口版.docx")
OUTPUT_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_字段增强版.docx")
CLUSTER_CSV = Path("/Users/dxk/xiaohongshu/data/daren_clusters_k3.csv")


FIELDS = [
    "活跃粉丝占比",
    "水粉占比",
    "粉丝女/男比例",
    "赞藏总数",
    "近60天平均点赞",
    "近60天平均收藏",
    "近60天平均评论",
    "近60天平均分享",
    "商业笔记总数",
    "图文笔记报价",
    "视频笔记报价",
    "图文CPE",
    "视频CPE",
]

ROW_LABELS = {
    "活跃粉丝占比": "活跃粉丝占比中位数",
    "水粉占比": "水粉占比中位数",
    "粉丝女/男比例": "粉丝女/男比例中位数",
    "赞藏总数": "赞藏总数中位数",
    "近60天平均点赞": "近60天平均点赞中位数",
    "近60天平均收藏": "近60天平均收藏中位数",
    "近60天平均评论": "近60天平均评论中位数",
    "近60天平均分享": "近60天平均分享中位数",
    "商业笔记总数": "商业笔记总数中位数",
    "图文笔记报价": "图文笔记报价中位数(元)",
    "视频笔记报价": "视频笔记报价中位数(元)",
    "图文CPE": "图文CPE中位数",
    "视频CPE": "视频CPE中位数",
}

GROUP_NAMES = {
    "0": "均衡主流型",
    "1": "收藏转化与商业合作型",
    "2": "高评论互动型",
}


def load_stats() -> tuple[dict[str, float], dict[str, dict[str, float]]]:
    with CLUSTER_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))

    overall = {}
    grouped = {g: {} for g in GROUP_NAMES}
    for field in FIELDS:
        vals = [float(r[field]) for r in rows if r.get(field) not in ("", None)]
        overall[field] = statistics.median(vals)
        for g in GROUP_NAMES:
            gvals = [float(r[field]) for r in rows if r.get("群体标签_k3") == g and r.get(field) not in ("", None)]
            grouped[g][field] = statistics.median(gvals)
    return overall, grouped


def fmt(field: str, value: float) -> str:
    if field in {"活跃粉丝占比", "水粉占比"}:
        return f"{value * 100:.2f}%"
    if field == "粉丝女/男比例":
        return f"{value:.1f}"
    if field in {"图文CPE", "视频CPE"}:
        return f"{value:.2f}"
    if field in {"图文笔记报价", "视频笔记报价", "赞藏总数", "近60天平均点赞", "近60天平均收藏", "近60天平均评论", "近60天平均分享", "商业笔记总数"}:
        return f"{int(round(value))}"
    return str(value)


def find_paragraph(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(prefix)


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style is not None:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        if paragraph.runs:
            src = paragraph.runs[0]
            run.bold = src.bold
            run.italic = src.italic
            run.underline = src.underline
            if src.font.name:
                run.font.name = src.font.name
                try:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), src.font.name)
                except Exception:
                    pass
            if src.font.size:
                run.font.size = src.font.size
    return new_para


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_table_three_line(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ["left", "right", "insideH", "insideV"]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")
    for edge, size in [("top", "12"), ("bottom", "12")]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")

    # header bottom line
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"})


def main():
    overall, grouped = load_stats()
    doc = Document(str(INPUT_DOC))

    # Update Chinese abstract
    p = find_paragraph(doc, "本文以小红书平台3403位粉丝数10万以上的生活记录类达人为研究对象")
    p.text = (
        "本文以小红书平台3403位粉丝数10万以上的生活记录类达人为研究对象，综合运用描述性统计分析、文本挖掘（帖子关键词共现、评论互动类型归纳与群体对比）及K-means聚类算法，围绕三个研究问题展开分析。"
        "研究发现如下：第一，该群体整体呈现女性主导（66.62%）、腰部达人为主（87.4%）、独立运营居多（61.74%）的结构特征，地域上高度集中于高线省市，核心受众为18至34岁年轻女性，商业笔记占比中位数仅为4.7%。同时，活跃粉丝占比中位数为39.75%，水粉占比中位数仅为7.52%，说明高影响力达人样本整体具备较稳定的受众质量。"
        "第二，高影响力达人在内容策略上普遍采用“泛生活+垂直兴趣”的复合布局，并表现出以“超级分享者”为核心的人设特征；评论文本显示，颜值赞美、拟亲密互动、提问求回应和消费转化构成了主要互动类型。"
        "第三，经肘部法则与轮廓系数验证，K-means聚类将达人生态划分为均衡主流型、收藏转化与商业合作型、高评论互动型三类群体，各群体不仅在粉丝画像和内容表达上存在差异，也在近60天互动质量与商业化表现上呈现出不同特征。"
    )

    # Update English abstract
    p = find_paragraph(doc, "The rise of content-driven e-commerce has reshaped both consumer decision-making")
    p.text = (
        "The rise of content-driven e-commerce has reshaped both consumer decision-making and brand communication. On platforms such as Xiaohongshu, high-influence lifestyle creators occupy a central position in this process, yet systematic empirical descriptions of their profile structure, content expression, and group differentiation remain limited. "
        "This study examines 3,403 Xiaohongshu lifestyle creators with at least 100,000 followers, and combines descriptive statistics, text mining of posts and comments, and K-means clustering to describe their behavioral patterns. The findings are threefold. First, the creator ecosystem is structurally concentrated: female creators account for 66.62% of the sample, mid-tier creators make up 87.4%, and 61.74% operate independently. Their core audience is concentrated among women aged 18–34 in higher-tier cities, while fan-quality indicators such as active-fan share and low suspicious-fan share remain relatively stable overall. "
        "Second, high-influence creators tend to adopt a hybrid content strategy that combines broad lifestyle expression with selected vertical interests, while maintaining closeness through everyday persona construction and comment interaction. Post and comment texts show that appearance praise, parasocial intimacy, question-response interaction, and light consumption conversion form the main interaction patterns. "
        "Third, K-means clustering identifies three distinct groups: balanced mainstream creators, collection-and-commercial-conversion creators, and high-comment-interaction creators. These groups differ not only in follower structure and content expression, but also in engagement quality and commercialization performance. Overall, the study provides a descriptive account of how user profiles, text expression, and creator differentiation are linked on Xiaohongshu."
    )

    # Update data counts in body + table 4-1
    p = find_paragraph(doc, "在第三章已完成研究对象界定、变量定义与方法路径说明的基础上")
    p.text = (
        "在第三章已完成研究对象界定、变量定义与方法路径说明的基础上，本节仅对第四章实际使用的数据口径作简要交代。第四章所用达人画像样本为3403位达人、聚类样本为3371位达人、文本样本为9731条帖子与123357条评论，对齐后的群体比较样本为9711条帖子与123176条评论，具体见表4-1。"
    )
    p = find_paragraph(doc, "在互动文本层面，本文进一步对")
    p.text = (
        "在互动文本层面，本文进一步对123176条已对齐评论进行归纳分析。为提高评论分析的可解释性，本文依据高频表达、语义模式与互动意图，将评论划分为颜值赞美、拟亲密互动、关怀支持、提问求回应、消费转化等类别，具体识别规则见表4-2。例如，“好美”“绝美”“好看”等表达被归入颜值赞美，“老婆”“宝宝”“想你”等表达被归入拟亲密互动，“怎么”“什么时候”“求回复”等表达被归入提问求回应，“同款”“链接”“下单”等表达则归入消费转化。结果显示，三类群体的评论都以泛互动类表达为主，但仍存在明显差异：群体0的颜值赞美与拟亲密互动相对更多，群体1更常出现提问求回应与消费转化表达，群体2整体评论结构更分散。"
    )
    p = find_paragraph(doc, "本研究以小红书平台3403位粉丝数10万以上的生活记录类达人为样本")
    p.text = (
        "本研究以小红书平台3403位粉丝数10万以上的生活记录类达人为样本，并结合9731条帖子与123357条评论的文本数据，围绕三个研究问题展开系统分析。以下结论分别对应达人画像特征、内容与互动特征以及达人群体分化三个层面。"
    )

    # Update table 4-1
    t1 = doc.tables[0]
    t1.rows[1].cells[1].text = "9731"
    t1.rows[1].cells[2].text = "9731"
    t1.rows[1].cells[3].text = "1901"
    t1.rows[2].cells[1].text = "123357"
    t1.rows[2].cells[2].text = "9001"
    t1.rows[2].cells[3].text = "1814"
    t1.rows[3].cells[1].text = "9711"
    t1.rows[3].cells[2].text = "9711"
    t1.rows[3].cells[3].text = "1897"
    t1.rows[4].cells[1].text = "123176"
    t1.rows[4].cells[2].text = "8986"
    t1.rows[4].cells[3].text = "1811"

    # Strengthen 4.2
    p = find_paragraph(doc, "本节首先从粉丝画像的视角对达人群体进行宏观描述")
    p.text = (
        "本节首先从粉丝画像的视角对达人群体进行宏观描述，考察用户画像特征（性别、年龄、地域、活跃时段）如何勾勒达人的受众基本盘。描述性统计结果显示，生活记录类达人以女性和独立创作者为主，分别约占样本的66%和61.7%。地域上高度集中于广东、浙江、上海、北京等高线省市，受众则主要分布于18至34岁的年轻女性群体。整体来看，该群体在商业化上较为克制，商业笔记占比中位数仅为4.7%；在互动结构上，收藏行为较为突出，藏赞比中位数为0.13，说明内容的保存价值在平台互动中占有重要位置。进一步比较不同粉丝画像的达人可以发现，面向25至34岁用户的达人更常表现出较高的收藏倾向，而面向更年轻用户的达人更容易引发评论互动。这一差异表明，不同年龄圈层的受众会以不同方式表达对达人内容的认同。"
    )
    anchor = p
    p2 = insert_paragraph_after(
        anchor,
        "从粉丝质量与互动质量看，高影响力达人样本整体表现出较为稳定的受众基础：活跃粉丝占比中位数为39.75%，水粉占比中位数为7.52%，粉丝女/男比例中位数为11.7，说明其受众不仅具有较高的女性集中度，也具有相对稳定的活跃度。与此同时，赞藏总数中位数为1378188，近60天平均点赞、收藏、评论和分享中位数分别为972、138、41和32，表明该群体的高影响力并不只体现在粉丝规模上，也体现在持续的互动质量和内容响应能力上。",
        style=p.style,
    )
    insert_paragraph_after(
        p2,
        "从商业化表现看，商业笔记总数中位数为10，图文与视频两类内容的报价中位数分别为15000元和20000元，视频笔记报价整体高于图文笔记报价；对应的图文CPE与视频CPE中位数分别为0和7.07，说明不同内容形式在商业承接与互动成本上存在差异。综合来看，小红书高影响力生活类达人不仅具备明确的受众画像特征，也在粉丝质量、互动质量与商业化能力层面表现出相对稳定的结构性特征。",
        style=p.style,
    )

    # Strengthen 4.4.2 and insert new table
    p = find_paragraph(doc, "基于K-means聚类结果，本文将3371位高影响力达人划分为三类特征鲜明的群体")
    p.text = (
        "基于K-means聚类结果，本文将3371位高影响力达人划分为三类特征鲜明的群体，相关画像特征概览见表4-4。群体0共有2671位达人，占比最高，整体表现较为均衡，可概括为“均衡主流型”；群体1共有617位达人，在藏赞比和商业笔记占比上明显高于整体均值，可概括为“收藏转化与商业合作型”；群体2共有83位达人，评赞比显著高于其余两类，可概括为“高评论互动型”。三类群体的并存表明，小红书高影响力达人生态并非均质结构，而是存在明显的功能分化。进一步结合粉丝质量、互动质量与商业化指标，可以更清楚地识别三类群体在受众基础和变现路径上的差异，详见表4-5。"
    )
    p = find_paragraph(doc, "进一步结合文本挖掘结果可以发现，不同群体在内容与互动方式上呈现出较为清晰的差异")
    p.text = (
        "进一步结合文本挖掘结果和增强指标可以发现，不同群体在内容、互动和商业化方式上呈现出较为清晰的差异。均衡主流型在活跃粉丝占比、水粉占比和近60天互动质量上整体更接近总体中位水平，表现出较稳定的受众基础和内容承接能力；收藏转化与商业合作型在商业笔记总数、收藏导向互动以及报价水平上更为突出，显示出较强的内容保存价值与商业承接能力；高评论互动型则在评论导向互动上更为明显，但其商业笔记总数和报价水平相对较低，说明互动热度与商业合作强度并不完全同步。总体来看，达人画像变量、内容策略与评论互动之间呈现出较为稳定的对应关系，共同构成了高影响力达人差异化发展的外在表现。从理论回应看，均衡主流型可从意见领袖理论中依靠持续内容供给维持广泛影响力的路径进行理解；收藏转化与商业合作型则表现出更强的内容保存价值与商业承接能力；高评论互动型可从社会认同理论视角理解，其较强的评论参与和关系表达更接近以互动认同维系影响力的路径。换言之，三类群体并非简单的流量高低差异，而是对应了不同的受众连接方式与影响力实现路径。"
    )

    # Insert caption and new three-line table after the analytical paragraph
    caption_ref = p
    caption = insert_paragraph_after(caption_ref, "表4-5 三类群体粉丝质量、互动质量与商业化特征比较", style=find_paragraph(doc, "表4-4 三类群体画像特征概览").style)
    table = doc.add_table(rows=1, cols=4)
    table._tbl.getparent().remove(table._tbl)
    caption._p.addnext(table._tbl)
    set_table_three_line(table)

    header = ["指标", "均衡主流型", "收藏转化与商业合作型", "高评论互动型"]
    for j, text in enumerate(header):
        set_cell_text(table.rows[0].cells[j], text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for field in FIELDS:
        row = table.add_row().cells
        set_cell_text(row[0], ROW_LABELS[field], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[1], fmt(field, grouped["0"][field]))
        set_cell_text(row[2], fmt(field, grouped["1"][field]))
        set_cell_text(row[3], fmt(field, grouped["2"][field]))

    # Fifth chapter sync
    p = find_paragraph(doc, "第一，在达人画像特征层面")
    p.text = (
        "第一，在达人画像特征层面，高影响力生活类达人整体呈现女性主导（66.62%）、腰部达人为主（87.4%）、独立运营居多（61.74%）的结构特征，地域上主要集中于浙江、广东、上海、北京等高线省市，核心受众为18至34岁的年轻女性群体。除粉丝规模外，该群体在粉丝质量和互动质量上也表现出相对稳定的特征：活跃粉丝占比中位数为39.75%，水粉占比中位数仅为7.52%，近60天平均点赞、收藏、评论和分享中位数分别为972、138、41和32。"
    )
    p = find_paragraph(doc, "第三，在达人群体分化层面")
    p.text = (
        "第三，在达人群体分化层面，经肘部法则与轮廓系数验证，K=3是当前样本下较为合适的聚类方案。均衡主流型（2671人）在各项指标上整体较为平衡，是平台中的主流高影响力达人；收藏转化与商业合作型（617人）在收藏导向互动、商业笔记总数和商业合作承接能力上表现更突出；高评论互动型（83人）则表现出更强的评论互动倾向，但其商业合作水平并未同步提升。三类群体的并存说明，小红书高影响力达人并不存在单一成功路径，而是表现出多样化的受众结构、内容表达、互动质量与商业化组合方式。"
    )
    p = find_paragraph(doc, "从理论层面看，本文在意见领袖理论和社会认同理论的基础上")
    p.text = (
        "从理论层面看，本文在意见领袖理论和社会认同理论的基础上，进一步提供了高影响力达人差异化路径的描述性证据：不同达人并非仅在流量规模上存在差异，更在受众结构、粉丝质量、互动质量和商业化表现的组合上展现出不同的影响力实现方式。从实践层面看，研究结果提示达人自身在内容定位之外也需要关注受众质量与互动结构，品牌方在达人筛选时不宜只看粉丝规模，平台则可基于创作者在互动质量和商业承接能力上的差异开展更精细的创作者分层。"
    )

    doc.save(str(OUTPUT_DOC))
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
