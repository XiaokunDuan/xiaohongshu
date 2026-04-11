#!/usr/bin/env python3
"""
Final visual polish: split a few long paragraphs, tune caption spacing,
and rebalance table 4-5 widths.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


INPUT_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_段落优化终版.docx")
OUTPUT_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_版式微调终版.docx")


def find_para(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(prefix)


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    new_para.style = paragraph.style
    run = new_para.add_run(text)
    if paragraph.runs:
        src = paragraph.runs[0]
        run.bold = src.bold
        run.italic = src.italic
        run.underline = src.underline
        if src.font.name:
            run.font.name = src.font.name
            try:
                run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", src.font.name)
            except Exception:
                pass
        if src.font.size:
            run.font.size = src.font.size
    return new_para


def split_para(paragraph, marker: str):
    text = paragraph.text
    if marker not in text:
        return
    left, right = text.split(marker, 1)
    paragraph.text = left.rstrip()
    insert_after(paragraph, marker + right)


def set_caption_spacing(paragraph, before=6, after=3):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def set_body_spacing(paragraph, after=6):
    paragraph.paragraph_format.space_after = Pt(after)


def main():
    doc = Document(str(INPUT_DOC))

    # Split a few remaining long body paragraphs
    split_para(find_para(doc, "在用户画像主线下，文本挖掘的作用并不是脱离画像单独讨论内容"), "前者主要对应")
    split_para(find_para(doc, "结合图1和图2可以进一步发现"), "因此，")
    split_para(find_para(doc, "图1呈现了达人标签的共现网络结构"), "可以观察到，")
    split_para(find_para(doc, "内容策略不仅体现在发什么，还体现在何时发"), "从星期维度看，")
    split_para(find_para(doc, "数据中同时包含图文和视频两套报价与互动成本数据"), "从互动成本（CPE）看，")
    split_para(find_para(doc, "进一步结合文本挖掘结果和增强指标可以发现"), "收藏转化与商业合作型")
    split_para(find_para(doc, "总体来看，达人画像变量、内容策略与评论互动之间呈现出较为稳定的对应关系"), "从理论回应看，")

    # Tune caption spacing around chapter 4 figures/tables
    for prefix in [
        "图1 达人标签共现词网络",
        "图2 高影响力达人粉丝关注焦点分布",
        "图3 高影响力达人简介高频词",
        "图4 高影响力达人粉丝高峰活跃时段与星期分布",
        "图5 图文与视频内容形式效能对比",
        "图6 肘部法则与轮廓系数K值选取",
        "表4-1 数据来源与样本规模",
        "表4-2 文本预处理与评论互动识别规则",
        "表4-3 聚类变量定义与入模理由",
        "表4-4 三类群体画像特征概览",
        "表4-5 三类群体粉丝质量、互动质量与商业化特征比较",
    ]:
        try:
            p = find_para(doc, prefix)
        except ValueError:
            continue
        set_caption_spacing(p, before=8, after=4)

    # Give surrounding analytical paragraphs a bit more breathing room
    for prefix in [
        "基于K-means聚类结果，本文将3371位高影响力达人划分为三类特征鲜明的群体",
        "三类群体的并存表明，小红书高影响力达人生态并非均质结构",
        "进一步结合文本挖掘结果和增强指标可以发现",
        "总体来看，达人画像变量、内容策略与评论互动之间呈现出较为稳定的对应关系",
    ]:
        try:
            set_body_spacing(find_para(doc, prefix), after=6)
        except ValueError:
            pass

    # Rebalance table 4-5 widths
    table5 = doc.tables[4]
    widths = [Cm(7.0), Cm(3.0), Cm(4.5), Cm(3.5)]
    for row in table5.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    doc.save(str(OUTPUT_DOC))
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
