from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


SRC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_五章最终版.docx")
OUT = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_五章送审终版.docx")


def find_paragraph(doc, startswith: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(startswith)


def insert_paragraph_before(paragraph: Paragraph, text: str, style) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.text = text
    return new_para


def insert_paragraph_after(paragraph: Paragraph, text: str, style) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.text = text
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def remove_all_between(doc, start_text: str, end_text: str) -> None:
    body = doc.element.body
    blocks = list(body.iterchildren())

    def idx(text: str) -> int:
        for i, child in enumerate(blocks):
            if child.tag == qn("w:p"):
                p = Paragraph(child, doc)
                if p.text.strip().startswith(text):
                    return i
        raise ValueError(text)

    s = idx(start_text)
    e = idx(end_text)
    for el in blocks[s + 1:e]:
        body.remove(el)


def main():
    doc = Document(str(SRC))
    chapter_style = find_paragraph(doc, "第五章 讨论与结论").style
    section_style = find_paragraph(doc, "4.3 高影响力达人行为模式的文本表现").style
    body_style = find_paragraph(doc, "本文以小红书平台中的高影响力生活记录类达人为研究对象").style

    # Repair chapter 3 heading and remove stray captions.
    p31 = find_paragraph(doc, "3.1 研究对象与数据来源")
    insert_paragraph_before(p31, "第三章 研究设计", chapter_style)
    seen_counts = {
        "表3-1 数据来源与样本规模": 0,
        "表3-3 文本预处理与评论互动识别规则": 0,
    }
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t == "表3-3 聚类变量定义与入模理由":
            remove_paragraph(p)
        elif t in seen_counts:
            seen_counts[t] += 1
            if t == "表3-1 数据来源与样本规模" and seen_counts[t] == 1:
                remove_paragraph(p)
            elif t == "表3-3 文本预处理与评论互动识别规则" and seen_counts[t] > 1:
                remove_paragraph(p)

    # Rebuild chapter 5 cleanly.
    remove_all_between(doc, "第五章 讨论与结论", "参考文献")
    ch5 = find_paragraph(doc, "第五章 讨论与结论")
    p = insert_paragraph_after(ch5, "5.1 主要研究发现", section_style)
    p = insert_paragraph_after(p, "本研究以小红书平台3403位粉丝数10万以上的生活记录类达人为样本，并结合9731条帖子与123357条评论的文本数据，围绕一个核心研究问题展开分析，即在用户画像视角下，高影响力生活记录类达人呈现出怎样的行为模式，以及这些行为模式如何在内容表达、互动结构与商业化表现上分化。", body_style)
    p = insert_paragraph_after(p, "首先，从画像结构看，高影响力生活类达人整体呈现女性主导、腰部达人为主、独立运营居多的特征，核心受众主要集中于18至34岁的年轻女性群体。除粉丝规模外，该群体在活跃粉丝占比、水粉占比以及近60天平均互动指标上也表现出相对稳定的结构性特征，说明高影响力并不仅体现在规模上，也体现在受众质量与持续互动能力上。", body_style)
    p = insert_paragraph_after(p, "其次，从群体划分看，经肘部法则与轮廓系数综合判断，K=3是当前样本下较为合适的聚类方案。均衡主流型在各项指标上整体较为平衡，是平台中的主流高影响力达人；收藏转化与商业合作型在收藏导向互动、商业笔记总数和商业合作承接能力上表现更突出；高评论互动型则表现出更强的评论互动倾向，但其商业合作水平并未同步提升。", body_style)
    p = insert_paragraph_after(p, "最后，从文本表现看，高影响力达人普遍采用“泛生活+垂直兴趣”的复合内容布局，并通过日常化、可亲近的人设表达维系与粉丝关系。帖子文本呈现出明显的生活方式化、情绪氛围化与场景化特征；评论文本则显示，颜值赞美、拟亲密互动、提问求回应与轻度消费转化构成主要互动类型。结合群体比较可以看到，行为模式并非单一路径，而是表现出多样化的受众结构、内容表达、互动质量与商业化组合方式。", body_style)

    p = insert_paragraph_after(p, "5.2 理论讨论", section_style)
    p = insert_paragraph_after(p, "从理论层面看，本文围绕“高影响力达人呈现出怎样的行为模式及其分化”这一核心问题，补充了对意见领袖理论和社会认同理论的描述性理解。意见领袖理论通常强调具有较高可见度和说服力的节点在传播链路中的作用，而本文发现高影响力达人并不存在单一路径：均衡主流型更接近依靠持续内容供给维持广泛影响力的典型路径，收藏转化与商业合作型则表现出更强的保存价值与商业承接能力，高评论互动型则更多依赖高参与评论和关系表达来维系影响力。社会认同理论则有助于理解评论区中的拟亲密互动、颜值赞美与求回应表达。本文发现，评论互动差异虽然具有统计显著性，但效应量总体较弱，这表明社会认同并非只在单一群体中出现，而是以不同强度分布在多个高影响力达人群体之中。", body_style)

    p = insert_paragraph_after(p, "5.3 文献回应与研究贡献", section_style)
    p = insert_paragraph_after(p, "与第二章文献述评相比，本文的推进主要体现在三个方面：其一，在研究对象上，从普通用户、单次营销效果或单条内容分析转向高影响力达人供给侧，补充了高影响力创作者群体本身的系统画像；其二，在分析框架上，将用户画像、群体划分与文本表现纳入同一解释路径，避免将达人视为同质化的营销节点；其三，在证据形式上，通过结构化画像数据与帖子、评论文本的对齐分析，进一步展示了不同类型达人在内容表达、互动结构与商业化承接上的差异，从而为现有文献提供了一种更接近创作者分层现实的描述性证据。", body_style)

    p = insert_paragraph_after(p, "5.4 实践启示", section_style)
    p = insert_paragraph_after(p, "从实践层面看，达人自身不宜只围绕粉丝规模理解影响力，而应根据自身更接近的行为模式路径优化策略：均衡主流型更应保持稳定的泛生活内容供给与受众黏性，收藏转化与商业合作型可进一步强化内容的保存价值与商业承接能力，高评论互动型则需在维持互动热度的同时审慎推进商业合作。对于品牌方而言，达人筛选不宜只看粉丝规模和单一报价，而应结合投放目标进行匹配：若强调广泛触达，可优先考虑均衡主流型；若强调内容保存与消费转化，可优先考虑收藏转化与商业合作型；若强调互动热度与评论参与，则可考虑高评论互动型。对于平台而言，研究结果说明创作者分层不宜仅依据单一流量指标，而应结合互动质量、商业承接能力和受众结构开展更精细的治理与扶持。", body_style)

    p = insert_paragraph_after(p, "5.5 研究局限与未来展望", section_style)
    p = insert_paragraph_after(p, "尽管本文在数据规模与分析维度上具有一定系统性，但仍存在若干局限。首先，本文使用的是横截面数据，只能反映达人在特定时间节点的静态状态，无法捕捉其从起步到高影响力阶段的动态演变过程，也无法进一步识别变量变化与后续表现之间的因果关系。其次，研究对象聚焦于小红书单一平台，所得结论是否适用于抖音、快手、微博等其他内容平台，仍需结合不同平台的算法机制、用户结构与内容生态进行进一步验证。", body_style)
    insert_paragraph_after(p, "再次，结构化画像数据主要来自第三方平台整理结果，仍可能存在覆盖范围与口径选择上的平台偏差；评论与帖子文本虽然规模较大，但采集主要集中于论文研究阶段的单一时间窗口，仍受到抓取时点和可见范围的限制。最后，本文主要采用描述性统计、文本挖掘与聚类分析方法，并补充了轻量统计检验，但尚未进一步引入回归模型、人工标注一致性验证或更复杂的主题模型，因此当前结论更适合被理解为对高影响力达人群体特征及其差异的描述性证据。未来研究可在更长时间窗口下引入纵向数据，并结合更系统的文本模型与因果识别设计，进一步检验不同类型达人行为模式的稳定性与演化路径。", body_style)

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
