#!/usr/bin/env python3
"""
Split remaining long paragraphs after paragraph optimization.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


INPUT_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_段落优化版.docx")
OUTPUT_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_段落优化终版.docx")


def find_para(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(prefix)


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    new_para.style = paragraph.style
    run = new_para.add_run(text)
    if paragraph.runs:
        src = paragraph.runs[0]
        run.bold = src.bold
        run.italic = src.italic
        run.underline = src.underline
        if src.font.name:
            run.font.name = src.font.name
            try:
                run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", src.font.name)
            except Exception:
                pass
        if src.font.size:
            run.font.size = src.font.size


def split_para(paragraph, marker: str):
    text = paragraph.text
    left, right = text.split(marker, 1)
    paragraph.text = left.rstrip()
    insert_after(paragraph, marker + right)


def main():
    doc = Document(str(INPUT_DOC))
    split_para(find_para(doc, "基于K-means聚类结果，本文将3371位高影响力达人划分为三类特征鲜明的群体"), "三类群体的并存表明，")
    split_para(find_para(doc, "进一步结合文本挖掘结果和增强指标可以发现"), "总体来看，")
    doc.save(str(OUTPUT_DOC))
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
