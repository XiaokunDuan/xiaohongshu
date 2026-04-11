from copy import deepcopy
from docx import Document


SRC = "/Users/dxk/Downloads/段晓坤 毕业论文二稿_学术修订版_v2.docx"
OUT = "/Users/dxk/Downloads/段晓坤 毕业论文二稿_LDA补强版.docx"


def insert_paragraph_after(paragraph, text):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.paragraphs[
        paragraph._parent._element.index(new_p)
    ]
    if new_para.runs:
        for run in new_para.runs:
            run.text = ""
    else:
        new_para.add_run("")
    new_para.text = text
    if paragraph.style is not None:
        new_para.style = paragraph.style
    return new_para


doc = Document(SRC)

# 3.3 方法补充
for p in doc.paragraphs:
    if p.text.strip().startswith("在文本处理环节"):
        p.text = (
            "在文本处理环节，本文统一完成清洗、中文分词、停用词过滤、TF-IDF辅助筛词、"
            "高频词筛选、共现网络构建与评论互动规则识别。其中，共现关系以同一帖子中的关键词共现为统计窗口，"
            "评论互动类型识别采用基于关键词与语义模式的规则归类方式，以保证结果具有较好的可解释性与可复核性，"
            "相关规则见表3-3。TF-IDF辅助筛词和共词分析方法分别参考 Salton and Buckley（1988）与 "
            "Callon et al.（1983）的相关研究。考虑到本文重点在于对高影响力达人文本表现进行结构性描述，"
            "因此在方法选择上优先采用可解释性较强的文本分析路径。"
        )
        insert_paragraph_after(
            p,
            "为进一步验证帖子文本的主题聚合方向，本文另行采用 LDA 对帖子标题与正文进行了补充性主题提取，"
            "并在不同主题数方案中比较关键词可解释性。LDA结果不作为本文的主分析依据，"
            "而主要用于辅助判断高频词与共现网络所呈现的内容结构是否具有稳定的主题归纳基础。",
        )
        break

# 4.3.1 结果补充
for p in doc.paragraphs:
    if p.text.strip().startswith("从结果上看，帖子文本的高频词主要围绕"):
        insert_paragraph_after(
            p,
            "作为补充性检验，LDA 主题提取将帖子文本大体归纳为五类内容主题："
            "城市旅行与节庆体验、护肤抗老与成分种草、母婴营养与喂养场景、春日氛围与日常审美表达、"
            "以及情绪成长与工作叙事。该结果与前文基于高频词和共现网络得到的判断总体一致，"
            "说明高影响力生活类达人并非围绕单一赛道展开表达，而是在泛生活叙事中嵌入若干相对稳定的垂直主题。",
        )
        break

doc.save(OUT)
print(OUT)
