#!/usr/bin/env python3
"""
Optimize paragraph segmentation for thesis body text.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


INPUT_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_字段增强版.docx")
OUTPUT_DOC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_段落优化版.docx")


def find_para(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(prefix)


def delete_para(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    new_para.style = paragraph.style
    run = new_para.add_run(text)
    if paragraph.runs:
        src = paragraph.runs[0]
        run.bold = src.bold
        run.italic = src.italic
        run.underline = src.underline
        if src.font.name:
            run.font.name = src.font.name
            try:
                run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", src.font.name)
            except Exception:
                pass
        if src.font.size:
            run.font.size = src.font.size
    return new_para


def split_para(paragraph, marker: str):
    text = paragraph.text
    left, right = text.split(marker, 1)
    paragraph.text = left.rstrip()
    insert_after(paragraph, marker + right)


def main():
    doc = Document(str(INPUT_DOC))

    # Merge 1.3.1 research questions into one paragraph
    p69 = find_para(doc, "围绕研究目标，本文主要从以下三个层面展开分析。")
    p70 = find_para(doc, "第一，识别小红书高影响力生活记录类达人的整体画像特征")
    p71 = find_para(doc, "第二，分析高影响力达人在内容表达")
    p72 = find_para(doc, "第三，基于粉丝画像特征与互动指标")
    p69.text = (
        "围绕研究目标，本文主要从以下三个层面展开分析：第一，识别小红书高影响力生活记录类达人的整体画像特征，包括其基础属性、粉丝结构与商业化表现；"
        "第二，分析高影响力达人在内容表达、人设塑造与评论互动中的主要策略特征；第三，基于粉丝画像特征与互动指标，对达人群体进行聚类划分，并比较不同群体在内容策略与商业模式上的差异。"
    )
    delete_para(p70)
    delete_para(p71)
    delete_para(p72)

    # Merge the two summary paragraphs in 4.3.1
    p135 = find_para(doc, "结合图1和图2可以进一步发现")
    p136 = find_para(doc, "因此，内容生态层面的文本挖掘结果表明")
    p135.text = p135.text + p136.text
    delete_para(p136)

    # Split long paragraphs
    split_para(find_para(doc, "本节首先从粉丝画像的视角对达人群体进行宏观描述"), "整体来看，")
    split_para(find_para(doc, "在互动文本层面，本文进一步对123176条已对齐评论进行归纳分析"), "结果显示，")
    split_para(find_para(doc, "在用户画像视角下，聚类分析的目的不是单纯把达人分组"), "之所以未将")
    split_para(find_para(doc, "为确定最优聚类数K，本文综合使用肘部法则与轮廓系数"), "结果显示，")
    split_para(find_para(doc, "尽管本文在数据规模与分析维度上具有一定系统性"), "再次，")

    doc.save(str(OUTPUT_DOC))
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
