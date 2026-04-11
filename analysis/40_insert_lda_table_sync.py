from copy import deepcopy
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = "/Users/dxk/Downloads/段晓坤 毕业论文二稿_LDA主方法终版.docx"
OUT = "/Users/dxk/Downloads/段晓坤 毕业论文二稿_LDA主题表同步版.docx"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("left", "right", "top", "bottom"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_col_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    if p.style is not None:
        pass


def set_table_three_line(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ["left", "right", "insideH", "insideV"]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")
    for edge, size in [("top", "12"), ("bottom", "12")]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"})


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.paragraphs[paragraph._parent._element.index(new_p)]
    if style is not None:
        new_para.style = style
    new_para.text = text
    return new_para


def insert_table_after(paragraph, rows):
    parent = paragraph._parent
    table = parent.add_table(rows=1, cols=len(rows[0]), width=Pt(430))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, txt in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[j], txt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, txt in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[j], txt, bold=False, align=align)
    widths = [1800, 3300, 4700]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            set_col_width(cell, widths[j])
    set_table_three_line(table)
    paragraph._p.addnext(table._tbl)
    return table


doc = Document(SRC)
caption_style = None
for p in doc.paragraphs:
    if p.text.strip().startswith("表4-1 "):
        caption_style = p.style
        break

target_p = None
for p in doc.paragraphs:
    if p.text.strip().startswith("从结果上看，LDA 将帖子文本较为稳定地归纳为四类主题"):
        p.text = (
            "从结果上看，LDA 将帖子文本较为稳定地归纳为四类主题：一类以“人生、时间、工作、朋友”等词为代表，"
            "对应情绪成长与工作叙事；一类以“宝宝、营养、奶粉、妈妈”等词为代表，对应母婴营养与喂养场景；"
            "一类以“春天、温柔、拍照、春日”等词为代表，对应春日氛围与日常审美表达；还有一类以“皮肤、抗老、精华、护肤”等词为代表，"
            "对应护肤抗老与成分种草。具体主题名称、代表关键词与代表文本见表4-3a。按三类群体比较，群体0更集中于春日氛围、旅行体验和泛生活表达，"
            "群体1更常落在护肤种草与保存价值较强的内容上，群体2则相对更接近个人叙事和生活切片表达。"
        )
        target_p = p
    elif p.text.strip().startswith("在帖子主题提取阶段，本文对 K=3 至 K=7 的 LDA 方案进行了比较"):
        p.text = (
            "在帖子主题提取阶段，本文对 K=3 至 K=7 的 LDA 方案进行了比较。综合关键词可解释性、代表文本一致性与主题区分度看，"
            "K=4 的划分最为稳定，因此将其作为帖子文本分析的主结果，相关主题结果见表4-3a。四类主题分别表现为情绪成长与工作叙事、"
            "母婴营养与喂养场景、春日氛围与日常审美表达，以及护肤抗老与成分种草。"
        )
    elif p.text.strip().startswith("结合图4-3和图4-4可以进一步发现"):
        p.text = (
            "结合表4-3a、图4-3和图4-4可以进一步发现，LDA 主题结构与达人标签共现网络、粉丝关注焦点之间存在较高的一致性。"
            "无论从主题提取结果还是从标签共现关系看，小红书高影响力生活类达人都呈现出“泛生活底盘+垂直兴趣补充”的内容组织方式，"
            "这也为后续的人设建构与互动差异分析提供了文本基础。"
        )
    elif p.text.strip().startswith("最后，从文本表现看，高影响力达人普遍采用“泛生活+垂直兴趣”的复合内容布局"):
        p.text = (
            "最后，从文本表现看，高影响力达人普遍采用“泛生活+垂直兴趣”的复合内容布局。帖子层面的 LDA 主题提取显示，"
            "情绪成长、母婴营养、春日审美与护肤种草构成最稳定的四类内容主题（见表4-3a）；评论文本则显示，颜值赞美、拟亲密互动、"
            "提问求回应与轻度消费转化构成主要互动类型。结合群体比较可以看到，行为模式并非单一路径，而是表现出多样化的受众结构、"
            "内容主题、互动质量与商业化组合方式。"
        )

if target_p is None:
    raise RuntimeError("Could not find LDA result paragraph.")

caption = insert_paragraph_after(target_p, "表4-3a 帖子文本LDA主题提取结果", caption_style)
rows = [
    ["主题名称", "代表关键词", "代表文本"],
    ["情绪成长与工作叙事", "人生、时间、工作、朋友、世界、幸福", "“为自己鼓掌……跑步跑了这么久，要不跑个比赛吧……”"],
    ["母婴营养与喂养场景", "宝宝、营养、奶粉、妈妈、吸收、配方", "“四款热门奶粉测评，谁是断奶期三料冠军？”"],
    ["春日氛围与日常审美表达", "春天、温柔、拍照、春日、阳光、氛围", "“在流动生机里舒展自我……瞬间把人拉进温柔的春光里。”"],
    ["护肤抗老与成分种草", "皮肤、抗老、精华、护肤、修护、成分", "“这款雅诗兰黛白金面霜……抗老方面我真的没少做研究。”"],
]
insert_table_after(caption, rows)

doc.save(OUT)
print(OUT)
