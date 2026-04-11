from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


BASE = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_版式微调终版.docx")
OUT = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_单问题重构版.docx")


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            yield ("t", Table(child, doc))


def find_paragraph(doc, startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(f"Paragraph starting with {startswith!r} not found")


def paragraph_block_index(doc, startswith):
    for idx, (kind, obj) in enumerate(iter_block_items(doc)):
        if kind == "p" and obj.text.strip().startswith(startswith):
            return idx
    raise ValueError(f"Block for {startswith!r} not found")


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def move_block_range_before(doc, start_text, end_text, before_text):
    body = doc.element.body
    blocks = list(body.iterchildren())
    start = paragraph_block_index(doc, start_text)
    end = paragraph_block_index(doc, end_text)
    before = paragraph_block_index(doc, before_text)
    segment = blocks[start:end]
    for el in segment:
        body.remove(el)
    # recompute before index after removal
    before = paragraph_block_index(doc, before_text)
    for el in reversed(segment):
        body.insert(before, el)


def set_para_text(paragraph, text):
    paragraph.text = text


def set_para_normal(paragraph, text, style):
    paragraph.style = style
    paragraph.text = text


def append_reference(doc, text, style):
    ref_heading = find_paragraph(doc, "参考文献")
    # append just before final sectPr by inserting before the last body child
    new_p = OxmlElement("w:p")
    body = doc.element.body
    body.insert(len(body) - 1, new_p)
    para = Paragraph(new_p, doc)
    para.style = style
    para.text = text


def main():
    doc = Document(str(BASE))

    body_style = find_paragraph(doc, "围绕研究目标").style

    # Core structural move: cluster section before text section.
    move_block_range_before(
        doc,
        "4.4 基于 K-means 算法的达人群体聚类实现",
        "第五章 结论与展望",
        "4.3 基于文本挖掘的达人内容策略与人设特征解析",
    )

    # Chapter 1: single research question framing.
    set_para_text(find_paragraph(doc, "围绕研究目标"), "围绕本文的核心研究问题，研究将重点回答：基于用户画像视角，小红书高影响力生活记录类达人呈现出怎样的行为模式，这些行为模式在内容表达、互动结构与商业化表现上如何分化？")
    set_para_text(find_paragraph(doc, "本文以粉丝画像为核心切入变量"), "本文以粉丝画像为核心切入维度，并将“行为模式”界定为达人在特定受众结构约束下所呈现出的内容表达、互动方式与商业化特征的稳定组合。围绕这一核心问题，后文将从画像结构、群体划分与文本表现三个分析维度展开。")
    set_para_text(find_paragraph(doc, "本文主要采用描述性统计分析、文本挖掘与K-means聚类分析三种方法"), "本文主要采用描述性统计分析、文本挖掘与K-means聚类分析三种方法。描述性统计用于刻画高影响力生活记录类达人的基础属性、粉丝结构与商业化表现，以把握其行为模式的整体画像基础。")
    set_para_text(find_paragraph(doc, "文本挖掘主要用于处理达人简介、帖子标题与正文以及评论文本"), "文本挖掘主要用于处理达人简介、帖子标题与正文以及评论文本。具体包括统一清洗、中文分词、停用词过滤、TF-IDF辅助筛词、共现网络构建和评论互动类型归纳，以识别行为模式在内容表达与互动组织上的外显特征。")
    set_para_text(find_paragraph(doc, "K-means聚类分析则用于在用户画像视角下识别达人群体的差异化类型"), "K-means聚类分析则用于在用户画像视角下识别达人群体的差异化类型。本文在核心变量标准化基础上，对粉丝数、年龄结构、藏赞比、评赞比和商业笔记占比进行聚类，以比较不同群体在受众结构、内容策略与商业化表现上的组合差异。")
    set_para_text(find_paragraph(doc, "上述三种方法构成由“整体画像描述—文本特征识别—群体类型划分”递进展开的研究路径"), "上述三种方法共同服务于本文的一个核心研究问题，并从画像结构、群体划分和文本表现三个维度形成递进分析路径。与之对应的数据来源、变量设定和文本处理规则将在第三章中集中交代。")
    set_para_text(find_paragraph(doc, "本文的创新点主要体现在两个方面"), "本文的创新点主要体现在两个方面。其一，围绕一个核心研究问题，将用户画像、文本挖掘与群体划分纳入同一解释框架，以识别高影响力生活类达人的差异化行为模式。其二，在结构化画像数据之外，引入大规模帖子与评论文本，并将群体划分结果与文本表现联动分析，从而提升对高影响力达人内容供给、互动组织与商业化路径的综合解释能力。")
    set_para_text(find_paragraph(doc, "本文以小红书平台3403位粉丝数10万以上的生活记录类达人为研究对象，综合运用描述性统计分析"), "本文以小红书平台3403位粉丝数10万以上的生活记录类达人为研究对象，综合运用描述性统计分析、文本挖掘（帖子关键词共现、评论互动类型归纳与群体对比）及K-means聚类算法，围绕一个核心研究问题展开分析，即在用户画像视角下，高影响力达人呈现出怎样的行为模式，以及这些行为模式如何在内容表达、互动结构与商业化表现上分化。研究发现如下：第一，该群体整体呈现女性主导（66.62%）、腰部达人为主（87.4%）、独立运营居多（61.74%）的结构特征，地域上高度集中于高线省市，核心受众为18至34岁年轻女性，商业笔记占比中位数仅为4.7%。同时，活跃粉丝占比中位数为39.75%，水粉占比中位数仅为7.52%，说明高影响力达人样本整体具备较稳定的受众质量。第二，高影响力达人在内容策略上普遍采用“泛生活+垂直兴趣”的复合布局，并表现出以“超级分享者”为核心的人设特征；评论文本显示，颜值赞美、拟亲密互动、提问求回应和消费转化构成了主要互动类型。第三，经肘部法则与轮廓系数验证，K-means聚类将达人生态划分为均衡主流型、收藏转化与商业合作型、高评论互动型三类群体，各群体不仅在粉丝画像和内容表达上存在差异，也在近60天互动质量与商业化表现上呈现出不同特征。")

    # Chapter 2: align literature gaps and theory with one core question.
    set_para_text(find_paragraph(doc, "综合上述文献可以看到"), "综合上述文献可以看到，现有研究已经较充分地讨论了小红书平台的内容消费、营销逻辑与信任生成机制，但仍缺少一个同时结合用户画像、群体划分与文本表现的整合性框架来理解高影响力达人行为模式。现有研究要么停留于普通用户与单次营销效果，要么只关注单一文本层面或单一影响力指标，尚未系统说明高影响力达人在受众结构、内容表达、互动组织与商业化路径上的稳定差异。基于此，本文以高影响力生活类达人为研究对象，尝试通过多维画像、文本挖掘与聚类分析相结合的方式，对其行为模式进行描述性归纳。")
    set_para_text(find_paragraph(doc, "意见领袖理论为本文理解高影响力达人的传播作用提供了基础框架"), "意见领袖理论为本文理解高影响力达人的传播作用提供了基础框架。Katz and Lazarsfeld（1955）指出，在特定社群中具备较高可见度、较强说服力与持续内容供给能力的个体，往往能够在信息扩散与决策形成过程中发挥关键影响。对应到小红书平台，高影响力达人并不只是拥有较大的粉丝规模，更重要的是通过持续发布内容、维持互动和塑造可信形象来影响受众。因而，本文在群体划分和类型解释中关注粉丝规模、收藏评论结构、商业合作强度以及评论互动差异，目的在于从数据层面描述不同意见领袖路径的外在表现。")
    set_para_text(find_paragraph(doc, "社会认同理论进一步解释了受众为何会对不同达人形成持续关注"), "社会认同理论进一步解释了受众为何会对不同达人形成持续关注。Tajfel and Turner（1979）强调，个体会通过群体归属、身份投射与情感认同来建构自我，并在消费与互动过程中表达这种认同。对于小红书用户而言，关注达人、参与评论、收藏内容乃至购买同款，往往不仅是功能性选择，也包含对某种生活方式、审美取向和圈层身份的认同。因此，本文将粉丝年龄结构、关注焦点以及评论中的拟亲密互动、颜值赞美和求回应表达视为理解达人影响力差异的重要线索。")
    set_para_text(find_paragraph(doc, "基于上述两个理论，本文将研究问题拆解为三个相互衔接的层面"), "基于上述两个理论，本文围绕一个核心研究问题构建分析框架，并从三个相互衔接的维度展开：首先以用户画像特征刻画达人的受众结构基础；其次通过聚类分析识别不同类型达人的分化路径；最后借助帖子与评论文本观察各类行为模式在内容表达与互动结构上的外显表现。意见领袖理论主要用于理解群体划分及影响力路径差异，社会认同理论则用于理解评论互动中的参与、拟亲密和关系建构。")

    # New Chapter 3 heading + merged research design.
    set_para_text(find_paragraph(doc, "第三章 研究设计：指标体系、文本挖掘思路与聚类方法"), "第三章 研究设计与实证分析")
    set_para_text(find_paragraph(doc, "3.1 研究对象、样本范围与数据来源"), "3.1 研究设计与数据来源")
    set_para_text(find_paragraph(doc, "本文将研究对象界定为小红书平台中的高影响力生活记录类达人"), "本文以小红书平台中的高影响力生活记录类达人为研究对象，并以粉丝数达到10万作为进入样本池的操作性门槛。该标准主要用于界定研究范围，而非将高影响力简单等同于单一流量指标；后续分析仍结合互动表现、商业笔记占比与文本特征综合考察达人差异。结构化画像样本来自灰豚平台在论文研究阶段导出的达人总表，共3403位达人；经缺失值处理与异常值筛查后，最终用于K-means聚类的有效样本为3371位。")
    set_para_normal(find_paragraph(doc, "3.2 指标体系与核心变量定义"), "在指标体系上，本文从基础身份、互动表现、商业化水平和粉丝画像四个维度构建结构化变量，并将帖子标题、正文、简介与评论文本作为文本变量的主要来源。由于本文关注的是用户画像约束下的行为模式差异，因此用户画像与互动结构被置于解释框架的核心位置。", body_style)
    set_para_text(find_paragraph(doc, "为对高影响力达人群体进行立体化描述"), "在聚类变量选择上，本文优先保留能够同时反映账号体量、受众年龄结构、内容保存价值、讨论互动倾向与商业合作强度的指标，即粉丝数、Top1年龄段占比、藏赞比、评赞比和商业笔记占比。性别比例、地域分布、活跃粉丝占比、水粉占比以及报价、CPE、CPM等变量虽然具有解释价值，但更适合作为群体比较和结果解释变量，而不直接用于界定聚类边界。")
    set_para_text(find_paragraph(doc, "在聚类变量选择上，本文优先保留能够同时反映账号体量"), "本文所称“行为模式”，是指达人在特定受众结构约束下所呈现出的内容表达、互动方式与商业化表现的稳定组合。相较于泛泛讨论单一行为动作，本文更关注达人如何在相对稳定的受众基础上形成内容供给风格、评论互动路径与商业合作方式。")
    set_para_normal(find_paragraph(doc, "3.3 行为模式的操作性定义、分析框架与方法路径"), "在文本处理环节，本文统一完成清洗、中文分词、停用词过滤、TF-IDF辅助筛词、高频词筛选、共现网络构建与评论互动规则识别。其中，共现关系以同一帖子中的关键词共现为统计窗口，评论互动类型识别采用规则归类方式，以保证结果具有较好的可解释性与可复核性。TF-IDF辅助筛词和共词分析方法分别参考 Salton and Buckley（1988）与 Callon et al.（1983）的相关研究。", body_style)
    set_para_text(find_paragraph(doc, "本文所称“行为模式”，并非泛指达人的所有线上行为"), "在聚类实现上，本文使用 sklearn 的 KMeans 模型，初始化方式为 k-means++，设置 n_init=10、random_state=42，并参考 MacQueen（1967）以及 Arthur and Vassilvitskii（2007）关于 K-means 与初始化改进的讨论。剔除异常值后，再结合肘部法则和轮廓系数确定最终聚类数。")
    set_para_text(find_paragraph(doc, "基于这一界定，本文的方法路径可以概括为三步"), "本章首先展示高影响力达人的整体画像结构，其次完成群体划分与类型识别，最后再对不同群体的帖子与评论文本进行比较，由此回答本文关于高影响力达人行为模式及其分化的核心研究问题。")

    # Remove old chapter 4 heading and old 4.1 heading.
    remove_paragraph(find_paragraph(doc, "第四章 基于多维数据的达人行为特征与画像聚类分析"))
    remove_paragraph(find_paragraph(doc, "4.1 数据来源、获取与预处理"))
    set_para_text(find_paragraph(doc, "在第三章已完成研究对象界定、变量定义与方法路径说明的基础上"), "本章实际使用的达人画像样本为3403位达人、聚类样本为3371位达人、文本样本为9731条帖子与123357条评论，对齐后的群体比较样本为9711条帖子与123176条评论，具体见表3-1。")
    set_para_text(find_paragraph(doc, "在预处理环节，本文主要完成统一清洗"), "在数据对齐环节，本文将结构化样本与DOM抓取的帖子、评论数据按照creator_id进行映射，并在必要时进行群体标签回填，从而得到可用于群体比较的9711条帖子与123176条评论。")

    # Renumber and rename sections.
    set_para_text(find_paragraph(doc, "4.2 达人群体基础分布与宏观统计特征分析"), "3.2 高影响力达人画像结构分析")
    set_para_text(find_paragraph(doc, "4.4 基于 K-means 算法的达人群体聚类实现"), "3.3 基于 K-means 算法的达人群体划分与类型特征")
    set_para_text(find_paragraph(doc, "4.4.1 聚类特征变量的选择与模型设定"), "3.3.1 聚类特征变量的选择与模型设定")
    set_para_text(find_paragraph(doc, "4.4.2 三类高影响力达人群体的画像特征解析"), "3.3.2 三类高影响力达人群体的画像特征解析")
    set_para_text(find_paragraph(doc, "4.3 基于文本挖掘的达人内容策略与人设特征解析"), "3.4 高影响力达人行为模式的文本表现")
    set_para_text(find_paragraph(doc, "4.3.1 泛生活与垂直兴趣交织的内容生态网络解析"), "3.4.1 泛生活与垂直兴趣交织的内容生态网络解析")
    set_para_text(find_paragraph(doc, "4.3.2 信任机制下的人设与评论互动文本解码"), "3.4.2 信任机制下的人设与评论互动文本解码")
    set_para_text(find_paragraph(doc, "4.3.3 内容发布时机"), "3.4.3 内容发布时机")
    set_para_text(find_paragraph(doc, "4.3.4 内容形式选择"), "3.4.4 内容形式选择")
    set_para_text(find_paragraph(doc, "第五章 结论与展望"), "第四章 结论与展望")
    set_para_text(find_paragraph(doc, "5.1 研究主要结论"), "4.1 研究结论")
    set_para_text(find_paragraph(doc, "5.2 研究局限性"), "4.2 研究局限性")
    set_para_text(find_paragraph(doc, "5.3 理论与实践启示"), "4.3 理论与实践启示")

    # Update cluster section phrasing after reordering.
    set_para_text(find_paragraph(doc, "在用户画像视角下，聚类分析的目的不是单纯把达人分组"), "在用户画像视角下，聚类分析的目的不是单纯把达人分组，而是识别不同受众结构下达人在互动风格和商业化表现上的组合差异。参考 Chen（2024）关于小红书创作者相关性分析的聚类思路，并结合 MacQueen（1967）与 Arthur and Vassilvitskii（2007）对 K-means 及其初始化改进的讨论，本文选取粉丝数、Top1年龄段占比、藏赞比、评赞比和商业笔记占比五项指标构建 K-means 聚类模型。其中，粉丝数用于衡量账号的基础体量，Top1年龄段占比用于刻画粉丝年龄结构的集中程度，藏赞比和评赞比分别反映内容被保存与引发讨论的倾向，商业笔记占比则衡量达人的商业合作强度。")
    set_para_text(find_paragraph(doc, "之所以未将性别比例、地域分布、活跃粉丝占比"), "之所以未将性别比例、地域分布、活跃粉丝占比、水粉占比以及报价、CPE、CPM直接纳入聚类，是因为这些指标更适合作为群体解释和结果比较变量，而不直接用于界定聚类边界。聚类变量及其入模理由见表3-3。")
    set_para_text(find_paragraph(doc, "在异常值处理方面，本文重点检查了评赞比和商业笔记占比"), "在异常值处理方面，本文重点检查了评赞比和商业笔记占比等波动较大的指标，对极端异常样本进行剔除后，聚类样本由3403位达人收敛至3371位。剔除前后，各类群体的相对比例和核心特征保持总体稳定，说明本文的聚类结果并非由少数极端样本驱动。")
    set_para_text(find_paragraph(doc, "为确定最优聚类数K，本文综合使用肘部法则与轮廓系数"), "为确定最优聚类数K，本文综合使用肘部法则与轮廓系数，对K=2至K=8进行逐一比较。聚类实现采用 sklearn 的 KMeans，初始化方式为 k-means++，并设置 n_init=10、random_state=42，以降低随机初始化带来的扰动。图6给出了不同 K 值下 SSE 与轮廓系数的变化情况。")
    set_para_text(find_paragraph(doc, "结果显示，当K=3时，轮廓系数达到0.354"), "结果显示，当K=3时，轮廓系数达到0.354，为当前比较范围内的最高值，同时SSE曲线在K=2至K=3之间出现明显拐点。需要说明的是，0.354并不属于边界非常清晰的高分聚类结构，但对于生活记录类达人这类同质性较高、边界相对模糊的样本而言，仍可视为具有一定区分度并具备解释价值。综合判断，本文采用K=3的聚类方案对3371位达人进行群体划分。")
    set_para_text(find_paragraph(doc, "基于K-means聚类结果，本文将3371位高影响力达人划分为三类特征鲜明的群体"), "基于K-means聚类结果，本文将3371位高影响力达人划分为三类特征鲜明的群体，相关画像特征概览见表3-4。群体0共有2671位达人，占比最高，整体表现较为均衡，可概括为“均衡主流型”；群体1共有617位达人，在藏赞比和商业笔记占比上明显高于整体均值，可概括为“收藏转化与商业合作型”；群体2共有83位达人，评赞比显著高于其余两类，可概括为“高评论互动型”。")
    set_para_text(find_paragraph(doc, "三类群体的并存表明，小红书高影响力达人生态并非均质结构"), "三类群体的并存表明，小红书高影响力达人生态并非均质结构，而是存在明显的功能分化。进一步结合粉丝质量、互动质量与商业化指标，可以更清楚地识别三类群体在受众基础和变现路径上的差异，详见表3-5。")
    set_para_text(find_paragraph(doc, "进一步结合文本挖掘结果和增强指标可以发现"), "进一步结合增强指标可以发现，不同群体在互动质量和商业化方式上已呈现出较为清晰的差异。均衡主流型在活跃粉丝占比、水粉占比和近60天互动质量上整体更接近总体中位水平，表现出较稳定的受众基础和内容承接能力；")
    set_para_text(find_paragraph(doc, "总体来看，达人画像变量、内容策略与评论互动之间呈现出较为稳定的对应关系"), "总体来看，达人画像变量已经显示出较为稳定的分化结构；后文的文本比较将进一步说明，这些群体差异如何在内容策略与评论互动中得到外显。")
    set_para_text(find_paragraph(doc, "从理论回应看，均衡主流型可从意见领袖理论中依靠持续内容供给维持广泛影响力的路径进行理解"), "从理论解释上看，均衡主流型更接近意见领袖理论中依靠持续内容供给维持广泛影响力的路径，而高评论互动型则更接近以互动认同维系影响力的情形。三类群体并非简单的流量高低差异，而是对应了不同的受众连接方式与影响力实现路径。")

    # Update text section intro after the move.
    set_para_text(find_paragraph(doc, "在用户画像主线下，文本挖掘的作用并不是脱离画像单独讨论内容"), "在完成群体划分的基础上，本文进一步从帖子与评论文本出发，考察三类达人在内容表达、人设塑造与互动组织上的差异。文本挖掘并不脱离用户画像单独讨论内容，而是作为识别行为模式外显表现的工具。")
    set_para_text(find_paragraph(doc, "前者主要对应达人标签、帖子标题与正文"), "具体处理流程包括文本清洗、中文分词、停用词过滤、TF-IDF辅助筛词、高频词筛选、共现网络构建以及互动类型归纳，并按三类聚类群体完成映射比较。相关处理规则见表3-2。")
    set_para_text(find_paragraph(doc, "在互动文本层面，本文进一步对123176条已对齐评论进行归纳分析"), "在互动文本层面，本文进一步对123176条已对齐评论进行归纳分析。为提高评论分析的可解释性，本文依据高频表达、语义模式与互动意图，将评论划分为颜值赞美、拟亲密互动、关怀支持、提问求回应、消费转化等类别，具体识别规则见表3-2。例如，“好美”“绝美”“好看”等表达被归入颜值赞美，“老婆”“宝宝”“想你”等表达被归入拟亲密互动，“怎么”“什么时候”“求回复”等表达被归入提问求回应，“同款”“链接”“下单”等表达则被归入消费转化。")

    # Renumber captions and in-text references from chapter 4 to chapter 3.
    replacements = {
        "表4-1": "表3-1",
        "表4-2": "表3-2",
        "表4-3": "表3-3",
        "表4-4": "表3-4",
        "表4-5": "表3-5",
        "第四章": "第三章",
        "第五章": "第四章",
    }
    # Limit chapter replacements on already-updated headings by doing targeted cleanup later.
    for p in doc.paragraphs:
        text = p.text
        if not text:
            continue
        # avoid corrupting current chapter headings we already set
        if text.startswith("第四章 结论与展望") or text.startswith("第三章 研究设计与实证分析"):
            continue
        new = text
        for old, new_val in replacements.items():
            new = new.replace(old, new_val)
        if new != text:
            p.text = new

    # Restore intended chapter heading strings after broad replace.
    find_paragraph(doc, "第三章 研究设计与实证分析").text = "第三章 研究设计与实证分析"
    find_paragraph(doc, "第四章 结论与展望").text = "第四章 结论与展望"

    # Update conclusion for single question framing.
    set_para_text(find_paragraph(doc, "本研究以小红书平台3403位粉丝数10万以上的生活记录类达人为样本"), "本研究以小红书平台3403位粉丝数10万以上的生活记录类达人为样本，并结合9731条帖子与123357条评论的文本数据，围绕一个核心研究问题展开分析，即在用户画像视角下，高影响力生活记录类达人呈现出怎样的行为模式，以及这些行为模式如何在内容表达、互动结构与商业化表现上分化。")
    set_para_text(find_paragraph(doc, "第一，在达人画像特征层面"), "首先，从画像结构看，高影响力生活类达人整体呈现女性主导（66.62%）、腰部达人为主（87.4%）、独立运营居多（61.74%）的结构特征，地域上主要集中于浙江、广东、上海、北京等高线省市，核心受众为18至34岁的年轻女性群体。除粉丝规模外，该群体在粉丝质量和互动质量上也表现出相对稳定的特征：活跃粉丝占比中位数为39.75%，水粉占比中位数仅为7.52%，近60天平均点赞、收藏、评论和分享中位数分别为972、138、41和32。")
    set_para_text(find_paragraph(doc, "第二，在内容与互动特征层面"), "其次，从群体划分看，经肘部法则与轮廓系数验证，K=3是当前样本下较为合适的聚类方案。均衡主流型（2671人）在各项指标上整体较为平衡，是平台中的主流高影响力达人；收藏转化与商业合作型（617人）在收藏导向互动、商业笔记总数和商业合作承接能力上表现更突出；高评论互动型（83人）则表现出更强的评论互动倾向，但其商业合作水平并未同步提升。")
    set_para_text(find_paragraph(doc, "第三，在达人群体分化层面"), "最后，从文本表现看，高影响力达人普遍采用“泛生活+垂直兴趣”的复合内容布局，并通过日常化、可亲近的人设表达维系与粉丝关系。帖子文本呈现出明显的生活方式化、情绪氛围化与场景化特征；评论文本则显示，颜值赞美、拟亲密互动、提问求回应与轻度消费转化构成主要互动类型。结合群体比较可以看到，行为模式并非单一路径，而是表现出多样化的受众结构、内容表达、互动质量与商业化组合方式。")
    set_para_text(find_paragraph(doc, "从理论层面看，本文在意见领袖理论和社会认同理论的基础上"), "从理论层面看，本文围绕“高影响力达人呈现出怎样的行为模式及其分化”这一核心问题，补充了对意见领袖理论和社会认同理论的描述性理解。研究显示，不同达人并非仅在流量规模上存在差异，更在受众结构、粉丝质量、互动质量和商业化表现的组合上展现出不同的影响力实现方式：均衡主流型更接近依靠持续内容供给维持广泛影响力的意见领袖路径，高评论互动型则更接近以互动参与和关系认同维系影响力的情形。")

    # Add missing method/theory references if absent.
    ref_style = find_paragraph(doc, "陈明，尹嘉璐").style
    existing_refs = "\n".join(p.text for p in doc.paragraphs if p.text)
    new_refs = [
        "Katz, E. & Lazarsfeld, P. F. 1955. Personal Influence: The Part Played by People in the Flow of Mass Communications. Glencoe, IL: Free Press.",
        "Tajfel, H. & Turner, J. C. 1979. An integrative theory of intergroup conflict. In W. G. Austin & S. Worchel (eds.), The Social Psychology of Intergroup Relations, 33-47. Monterey, CA: Brooks/Cole.",
        "MacQueen, J. 1967. Some methods for classification and analysis of multivariate observations. In Proceedings of the Fifth Berkeley Symposium on Mathematical Statistics and Probability, 281-297.",
        "Arthur, D. & Vassilvitskii, S. 2007. k-means++: The advantages of careful seeding. In Proceedings of the Eighteenth Annual ACM-SIAM Symposium on Discrete Algorithms, 1027-1035.",
        "Salton, G. & Buckley, C. 1988. Term-weighting approaches in automatic text retrieval. Information Processing & Management, 24(5), 513-523.",
        "Callon, M., Courtial, J. P., Turner, W. A. & Bauin, S. 1983. From translations to problematic networks: An introduction to co-word analysis. Social Science Information, 22(2), 191-235.",
    ]
    for ref in new_refs:
        if ref.split(".")[0] not in existing_refs:
            append_reference(doc, ref, ref_style)

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
