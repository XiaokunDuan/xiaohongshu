from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


SRC = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_五章统计增强版.docx")
OUT = Path("/Users/dxk/Downloads/段晓坤 毕业论文二稿_五章最终版.docx")


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            yield ("t", Table(child, doc))


def paragraph_block_index(doc, startswith: str) -> int:
    for idx, (kind, obj) in enumerate(iter_block_items(doc)):
        if kind == "p" and obj.text.strip().startswith(startswith):
            return idx
    raise ValueError(startswith)


def find_paragraph(doc, startswith: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(startswith)


def remove_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def remove_between_headings(doc, start_text: str, end_text: str) -> None:
    body = doc.element.body
    blocks = list(body.iterchildren())
    start = paragraph_block_index(doc, start_text)
    end = paragraph_block_index(doc, end_text)
    for el in blocks[start + 1:end]:
        body.remove(el)


def remove_caption_and_table(doc, caption_text: str) -> None:
    body = doc.element.body
    blocks = list(body.iterchildren())
    idx = paragraph_block_index(doc, caption_text)
    segment = [blocks[idx]]
    if idx + 1 < len(blocks) and blocks[idx + 1].tag == qn("w:tbl"):
        segment.append(blocks[idx + 1])
    for el in segment:
        body.remove(el)


def move_caption_and_table_before(doc, caption_text: str, before_text: str) -> None:
    body = doc.element.body
    blocks = list(body.iterchildren())
    cap_idx = paragraph_block_index(doc, caption_text)
    segment = [blocks[cap_idx]]
    if cap_idx + 1 < len(blocks) and blocks[cap_idx + 1].tag == qn("w:tbl"):
        segment.append(blocks[cap_idx + 1])
    for el in segment:
        body.remove(el)
    before_idx = paragraph_block_index(doc, before_text)
    for el in reversed(segment):
        body.insert(before_idx, el)


def insert_paragraph_before(paragraph: Paragraph, text: str, style) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.text = text
    return new_para


def insert_paragraph_after(paragraph: Paragraph, text: str, style) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.text = text
    return new_para


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_table_three_line(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ["left", "right", "insideH", "insideV"]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")
    for edge, size in [("top", "12"), ("bottom", "12")]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"})


def insert_table_after(paragraph: Paragraph, rows: list[list[str]]):
    doc = paragraph._parent
    table = doc.add_table(rows=1, cols=len(rows[0]), width=Pt(430))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, text in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[j], text, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, text in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], text, bold=False, align=align)
    set_table_three_line(table)
    paragraph._p.addnext(table._tbl)
    return table


def build_chapter3(doc, chapter_style, section_style, body_style, caption_style):
    remove_between_headings(doc, "第三章 研究设计", "第四章 实证结果分析")

    chapter4 = find_paragraph(doc, "第四章 实证结果分析")
    insert_paragraph_before(chapter4, "第三章 研究设计", chapter_style)
    current = insert_paragraph_before(chapter4, "3.1 研究对象与数据来源", section_style)
    current = insert_paragraph_after(current, "本文以小红书平台中的高影响力生活记录类达人为研究对象，并以粉丝数达到10万作为进入样本池的操作性门槛。该标准主要用于界定研究范围，而非将高影响力简单等同于单一流量指标；后续分析仍结合互动表现、商业笔记占比与文本特征综合考察达人差异。结构化画像样本来自灰豚平台在论文研究阶段导出的达人总表，共3403位达人；经缺失值处理与异常值筛查后，最终用于K-means聚类的有效样本为3371位。", body_style)
    current = insert_paragraph_after(current, "本文实际使用的达人画像样本为3403位达人、聚类样本为3371位达人、文本样本为9731条帖子与123357条评论，对齐后的群体比较样本为9711条帖子与123176条评论，具体见表3-1。", body_style)
    current = insert_paragraph_after(current, "在数据对齐环节，本文将结构化样本与DOM抓取的帖子、评论数据按照 creator_id 进行映射，并在必要时进行群体标签回填，从而得到可用于群体比较的文本样本。", body_style)
    cap = insert_paragraph_after(current, "表3-1 数据来源与样本规模", caption_style)
    insert_table_after(cap, [
        ["数据层次", "样本量", "说明"],
        ["达人画像样本", "3403", "灰豚结构化总表"],
        ["聚类有效样本", "3371", "异常值处理后用于K-means"],
        ["帖子文本样本", "9731", "dom_crawl 主文本池"],
        ["评论文本样本", "123357", "dom_crawl 主文本池"],
        ["群体对齐帖子", "9711", "用于群体比较"],
        ["群体对齐评论", "123176", "用于群体比较"],
    ])

    current = insert_paragraph_after(cap, "3.2 指标体系与变量设定", section_style)
    current = insert_paragraph_after(current, "在指标体系上，本文从基础身份、互动表现、商业化水平和粉丝画像四个维度构建结构化变量，并将帖子标题、正文、简介与评论文本作为文本变量的主要来源。由于本文关注的是用户画像约束下的行为模式差异，因此用户画像与互动结构被置于解释框架的核心位置。", body_style)
    current = insert_paragraph_after(current, "本文所称“行为模式”，是指达人在特定受众结构约束下所呈现出的内容表达、互动方式与商业化表现的稳定组合。相较于泛泛讨论单一行为动作，本文更关注达人如何在相对稳定的受众基础上形成内容供给风格、评论互动路径与商业合作方式。", body_style)
    current = insert_paragraph_after(current, "在聚类变量选择上，本文优先保留能够同时反映账号体量、受众年龄结构、内容保存价值、讨论互动倾向与商业合作强度的指标，即粉丝数、Top1年龄段占比、藏赞比、评赞比和商业笔记占比。性别比例、地域分布、活跃粉丝占比、水粉占比以及报价、CPE、CPM等变量虽然具有解释价值，但更适合作为群体结果比较变量，而不直接作为决定分组边界的入模变量，相关定义见表3-2。", body_style)
    cap = insert_paragraph_after(current, "表3-2 聚类变量定义与入模理由", caption_style)
    insert_table_after(cap, [
        ["变量名称", "变量含义", "入模理由"],
        ["粉丝数", "达人账号的粉丝规模，用于衡量基础影响力", "反映达人受众规模差异"],
        ["Top1年龄段占比", "达人粉丝中占比最高年龄段的集中程度", "反映粉丝年龄结构的集中程度"],
        ["藏赞比", "近60天平均收藏与近60天平均点赞量之比", "反映内容被保存的倾向"],
        ["评赞比", "近60天平均评论与近60天平均点赞量之比", "反映内容引发讨论互动的能力"],
        ["商业笔记占比", "商业笔记总数占全部笔记总数的比例", "反映达人商业合作强度"],
    ])

    current = insert_paragraph_after(cap, "3.3 文本挖掘、聚类与统计检验设计", section_style)
    current = insert_paragraph_after(current, "在文本处理环节，本文统一完成清洗、中文分词、停用词过滤、TF-IDF辅助筛词、高频词筛选、共现网络构建与评论互动规则识别。其中，共现关系以同一帖子中的关键词共现为统计窗口，评论互动类型识别采用规则归类方式，以保证结果具有较好的可解释性与可复核性，相关规则见表3-3。TF-IDF辅助筛词和共词分析方法分别参考 Salton and Buckley（1988）与 Callon et al.（1983）的相关研究。", body_style)
    current = insert_paragraph_after(current, "在聚类实现上，本文使用 sklearn 的 KMeans 模型，初始化方式为 k-means++，设置 n_init=10、random_state=42，并参考 MacQueen（1967）以及 Arthur and Vassilvitskii（2007）关于 K-means 与初始化改进的讨论。剔除异常值后，再结合肘部法则和轮廓系数确定最终聚类数。", body_style)
    current = insert_paragraph_after(current, "为避免群体差异仅停留于描述性比较，本文进一步设置两类轻量统计检验：一是对群体与评论互动类型构成的列联表执行卡方检验，以判断互动结构差异是否具有统计意义；二是对关键连续变量执行 Kruskal-Wallis 检验，以识别不同群体在粉丝质量、互动质量与商业化指标上的中位数差异。", body_style)
    cap = insert_paragraph_after(current, "表3-3 文本预处理与评论互动识别规则", caption_style)
    insert_table_after(cap, [
        ["模块", "类别/步骤", "规则说明"],
        ["文本预处理", "文本清洗", "去除换行、URL与无意义空格，保留基础中文语义单元"],
        ["文本预处理", "中文分词", "使用 jieba 分词并统一词形"],
        ["文本预处理", "停用词过滤", "过滤助词、代词与平台噪声词"],
        ["文本预处理", "TF-IDF辅助筛词", "先按词频初筛，再结合TF-IDF剔除泛化词"],
        ["网络构建", "共现窗口与边权", "以同一帖子为共现窗口，共现次数作为边权"],
        ["互动识别", "评论互动规则", "依据关键词与语义模式识别颜值赞美、拟亲密互动、关怀支持、提问求回应、消费转化"],
    ])


def rebuild_chapter5(doc, section_style, body_style):
    body = doc.element.body
    blocks = list(body.iterchildren())
    start = paragraph_block_index(doc, "第五章 讨论与结论")
    end = paragraph_block_index(doc, "参考文献")
    for el in blocks[start + 1:end]:
        body.remove(el)

    chapter5 = find_paragraph(doc, "第五章 讨论与结论")
    p = insert_paragraph_after(chapter5, "5.1 主要研究发现", section_style)
    p = insert_paragraph_after(p, "本研究以小红书平台3403位粉丝数10万以上的生活记录类达人为样本，并结合9731条帖子与123357条评论的文本数据，围绕一个核心研究问题展开分析，即在用户画像视角下，高影响力生活记录类达人呈现出怎样的行为模式，以及这些行为模式如何在内容表达、互动结构与商业化表现上分化。", body_style)
    p = insert_paragraph_after(p, "首先，从画像结构看，高影响力生活类达人整体呈现女性主导、腰部达人为主、独立运营居多的特征，核心受众主要集中于18至34岁的年轻女性群体。除粉丝规模外，该群体在活跃粉丝占比、水粉占比以及近60天平均互动指标上也表现出相对稳定的结构性特征，说明高影响力并不仅体现在规模上，也体现在受众质量与持续互动能力上。", body_style)
    p = insert_paragraph_after(p, "其次，从群体划分看，经肘部法则与轮廓系数综合判断，K=3是当前样本下较为合适的聚类方案。均衡主流型在各项指标上整体较为平衡，是平台中的主流高影响力达人；收藏转化与商业合作型在收藏导向互动、商业笔记总数和商业合作承接能力上表现更突出；高评论互动型则表现出更强的评论互动倾向，但其商业合作水平并未同步提升。", body_style)
    p = insert_paragraph_after(p, "最后，从文本表现看，高影响力达人普遍采用“泛生活+垂直兴趣”的复合内容布局，并通过日常化、可亲近的人设表达维系与粉丝关系。帖子文本呈现出明显的生活方式化、情绪氛围化与场景化特征；评论文本则显示，颜值赞美、拟亲密互动、提问求回应与轻度消费转化构成主要互动类型。结合群体比较可以看到，行为模式并非单一路径，而是表现出多样化的受众结构、内容表达、互动质量与商业化组合方式。", body_style)

    p = insert_paragraph_after(p, "5.2 理论讨论", section_style)
    p = insert_paragraph_after(p, "从理论层面看，本文围绕“高影响力达人呈现出怎样的行为模式及其分化”这一核心问题，补充了对意见领袖理论和社会认同理论的描述性理解。意见领袖理论通常强调具有较高可见度和说服力的节点在传播链路中的作用，而本文发现高影响力达人并不存在单一路径：均衡主流型更接近依靠持续内容供给维持广泛影响力的典型路径，收藏转化与商业合作型则表现出更强的保存价值与商业承接能力，高评论互动型则更多依赖高参与评论和关系表达来维系影响力。社会认同理论则有助于理解评论区中的拟亲密互动、颜值赞美与求回应表达。本文发现，评论互动差异虽然具有统计显著性，但效应量总体较弱，这表明社会认同并非只在单一群体中出现，而是以不同强度分布在多个高影响力达人群体之中。", body_style)

    p = insert_paragraph_after(p, "5.3 文献回应与研究贡献", section_style)
    p = insert_paragraph_after(p, "与第二章文献述评相比，本文的推进主要体现在三个方面：其一，在研究对象上，从普通用户、单次营销效果或单条内容分析转向高影响力达人供给侧，补充了高影响力创作者群体本身的系统画像；其二，在分析框架上，将用户画像、群体划分与文本表现纳入同一解释路径，避免将达人视为同质化的营销节点；其三，在证据形式上，通过结构化画像数据与帖子、评论文本的对齐分析，进一步展示了不同类型达人在内容表达、互动结构与商业化承接上的差异，从而为现有文献提供了一种更接近创作者分层现实的描述性证据。", body_style)

    p = insert_paragraph_after(p, "5.4 实践启示", section_style)
    p = insert_paragraph_after(p, "从实践层面看，达人自身不宜只围绕粉丝规模理解影响力，而应根据自身更接近的行为模式路径优化策略：均衡主流型更应保持稳定的泛生活内容供给与受众黏性，收藏转化与商业合作型可进一步强化内容的保存价值与商业承接能力，高评论互动型则需在维持互动热度的同时审慎推进商业合作。对于品牌方而言，达人筛选不宜只看粉丝规模和单一报价，而应结合投放目标进行匹配：若强调广泛触达，可优先考虑均衡主流型；若强调内容保存与消费转化，可优先考虑收藏转化与商业合作型；若强调互动热度与评论参与，则可考虑高评论互动型。对于平台而言，研究结果说明创作者分层不宜仅依据单一流量指标，而应结合互动质量、商业承接能力和受众结构开展更精细的治理与扶持。", body_style)

    p = insert_paragraph_after(p, "5.5 研究局限与未来展望", section_style)
    p = insert_paragraph_after(p, "尽管本文在数据规模与分析维度上具有一定系统性，但仍存在若干局限。首先，本文使用的是横截面数据，只能反映达人在特定时间节点的静态状态，无法捕捉其从起步到高影响力阶段的动态演变过程，也无法进一步识别变量变化与后续表现之间的因果关系。其次，研究对象聚焦于小红书单一平台，所得结论是否适用于抖音、快手、微博等其他内容平台，仍需结合不同平台的算法机制、用户结构与内容生态进行进一步验证。", body_style)
    insert_paragraph_after(p, "再次，结构化画像数据主要来自第三方平台整理结果，仍可能存在覆盖范围与口径选择上的平台偏差；评论与帖子文本虽然规模较大，但采集主要集中于论文研究阶段的单一时间窗口，仍受到抓取时点和可见范围的限制。最后，本文主要采用描述性统计、文本挖掘与聚类分析方法，并补充了轻量统计检验，但尚未进一步引入回归模型、人工标注一致性验证或更复杂的主题模型，因此当前结论更适合被理解为对高影响力达人群体特征及其差异的描述性证据。未来研究可在更长时间窗口下引入纵向数据，并结合更系统的文本模型与因果识别设计，进一步检验不同类型达人行为模式的稳定性与演化路径。", body_style)


def main():
    doc = Document(str(SRC))
    chapter_style = find_paragraph(doc, "第五章 讨论与结论").style
    section_style = find_paragraph(doc, "5.1 主要研究发现").style
    body_style = find_paragraph(doc, "本研究以小红书平台3403位粉丝数10万以上的生活记录类达人为样本").style
    caption_style = find_paragraph(doc, "表4-1 三类群体画像特征概览").style

    # Remove stray design tables from chapter 4 before rebuilding chapter 3.
    remove_caption_and_table(doc, "表3-3 聚类变量定义与入模理由")
    remove_caption_and_table(doc, "表3-3 文本预处理与评论互动识别规则")
    remove_caption_and_table(doc, "表3-1 数据来源与样本规模")

    build_chapter3(doc, chapter_style, section_style, body_style, caption_style)

    # Place the enhanced comparison table inside 4.2 instead of chapter end.
    move_caption_and_table_before(doc, "表4-2 三类群体粉丝质量、互动质量与商业化特征比较", "4.3 高影响力达人行为模式的文本表现")
    for prefix in ["表3-3 聚类变量定义与入模理由", "表3-3 文本预处理与评论互动识别规则"]:
        try:
            remove_paragraph(find_paragraph(doc, prefix))
        except ValueError:
            pass

    rebuild_chapter5(doc, section_style, body_style)

    doc.save(str(OUT))


if __name__ == "__main__":
    main()
